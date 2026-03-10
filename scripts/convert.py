import argparse
import fnmatch
import logging
import os
import platform
import re
import shutil
import sys
import subprocess
import yaml
import zipfile
import xml.etree.ElementTree as ElTree
from defusedxml import ElementTree as DefusedElTree
from typing import Any, Dict, List, Tuple, cast
from operator import itemgetter
from itertools import groupby
from pathlib import Path
from pathvalidate.argparse import validate_filepath_arg
from pathvalidate import sanitize_filepath


class ConvertVars:
    BASE_PATH = os.path.split(os.path.dirname(os.path.realpath(__file__)))[0]
    EDITION_CHOICES: List[str] = ["all", "webapp", "mobileapp", "against-security"]
    FILETYPE_CHOICES: List[str] = ["all", "docx", "odt", "pdf", "idml"]
    LAYOUT_CHOICES: List[str] = ["all", "leaflet", "guide", "cards"]
    LANGUAGE_CHOICES: List[str] = ["all", "en", "es", "fr", "nl", "no-nb", "pt-pt", "pt-br", "hu", "it", "ru"]
    VERSION_CHOICES: List[str] = ["all", "latest", "1.0", "1.1", "2.2", "3.0", "5.0"]
    LATEST_VERSION_CHOICES: List[str] = ["1.1", "3.0"]
    TEMPLATE_CHOICES: List[str] = ["all", "bridge", "bridge_qr", "tarot", "tarot_qr"]
    EDITION_VERSION_MAP: Dict[str, Dict[str, str]] = {
        "webapp": {"2.2": "2.2", "3.0": "3.0"},
        "against-security": {"1.0": "1.0"},
        "mobileapp": {"1.0": "1.0", "1.1": "1.1"},
        "all": {"2.2": "2.2", "1.0": "1.0", "1.1": "1.1", "3.0": "3.0", "5.0": "5.0"},
    }
    DEFAULT_TEMPLATE_FILENAME: str = os.sep.join(
        ["resources", "templates", "owasp_cornucopia_edition_ver_layout_document_template_lang"]
    )
    DEFAULT_OUTPUT_FILENAME: str = os.sep.join(["output", "owasp_cornucopia_edition_ver_layout_document_template_lang"])
    args: argparse.Namespace
    can_convert_to_pdf: bool = False


def check_fix_file_extension(filename: str, file_type: str) -> str:
    if filename and not filename.endswith(file_type):
        filename_split = os.path.splitext(filename)
        if filename_split[1].strip(".").isnumeric():
            filename = filename + "." + file_type.strip(".")
        else:
            filename = ".".join([os.path.splitext(filename)[0], file_type.strip(".")])
        logging.debug(f" --- output_filename with new ext = {filename}")
    return filename


def check_make_list_into_text(var: List[str]) -> str:
    if not isinstance(var, list):
        return str(var)
    var = group_number_ranges(var)
    text_output = ", ".join(str(s) for s in var)
    if not text_output.strip():
        text_output = " - "

    return text_output


def _validate_file_paths(source_filename: str, output_pdf_filename: str) -> Tuple[bool, str, str]:
    """Validate and sanitize file paths to prevent command injection."""
    source_path = os.path.abspath(source_filename)
    output_dir = os.path.abspath(os.path.dirname(output_pdf_filename))

    # Additional security checks
    if not os.path.isfile(source_path):
        return False, f"Source file does not exist: {source_path}", ""

    if not os.path.isdir(output_dir):
        return False, f"Output directory does not exist: {output_dir}", ""

    # Ensure paths are within expected directories to prevent path traversal
    base_path = os.path.abspath(convert_vars.BASE_PATH)
    if not source_path.startswith(base_path):
        return False, f"Source path outside base directory: {source_path}", ""
    if not output_dir.startswith(base_path):
        return False, f"Output directory outside base directory: {output_dir}", ""

    return True, source_path, output_dir


def _safe_extractall(archive: zipfile.ZipFile, target_dir: str) -> None:
    """Extract zip members only if their resolved paths stay within target_dir.

    Prevents Zip Slip / path traversal (CWE-22) by resolving symlinks and all
    '..' components before comparing each member path against the target root.
    Degenerate root entries ('.', '', './') are skipped rather than extracted,
    because they carry no file content and resolve to the target directory itself.
    """
    abs_target = os.path.realpath(target_dir)
    for member in archive.infolist():
        member_path = os.path.realpath(os.path.join(abs_target, member.filename))

        # Root/degenerate entries ('.', '', './') resolve to abs_target itself.
        # They are directory metadata with no content; skip them safely.
        if member_path == abs_target:
            continue

        # Block any member whose resolved path escapes the target directory.
        # The os.sep suffix prevents prefix collisions (e.g. /tmp/d vs /tmp/d_evil).
        if not member_path.startswith(abs_target + os.sep):
            raise ValueError(f"Zip Slip blocked: member '{member.filename}' would extract outside target directory")

        archive.extract(member, target_dir)


def _validate_command_args(cmd_args: List[str]) -> bool:
    """Validate command arguments for dangerous characters."""
    dangerous_chars = ["&", "|", ";", "$", "`", "(", ")", "<", ">", "*", "?", "[", "]", "{", "}", "\\"]
    for arg in cmd_args:
        if any(char in arg for char in dangerous_chars):
            logging.warning(f"Potentially dangerous character found in argument: {arg}")
            return False
    return True


def _convert_with_libreoffice(source_filename: str, output_pdf_filename: str) -> bool:
    libreoffice_bin = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice_bin and platform.system() == "Windows":
        potential_soffice = Path("C:/Program Files/LibreOffice/program/soffice.exe")
        if potential_soffice.exists():
            libreoffice_bin = str(potential_soffice)

    if not libreoffice_bin:
        return False

    try:
        logging.info(f"Using LibreOffice for conversion: {libreoffice_bin}")

        # Validate file paths
        is_valid, source_path, output_dir = _validate_file_paths(source_filename, output_pdf_filename)
        if not is_valid:
            logging.warning(source_path)  # source_path contains the error message
            return False

        # Create user profile directory safely
        user_profile_dir = os.path.abspath(os.path.join(convert_vars.BASE_PATH, "output", "lo_profile"))
        os.makedirs(user_profile_dir, exist_ok=True)
        user_profile_url = "file:///" + user_profile_dir.replace("\\", "/")

        # Build command arguments
        cmd_args = [
            libreoffice_bin,
            "--headless",
            f"-env:UserInstallation={user_profile_url}",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            source_path,
        ]

        # Validate command arguments
        if not _validate_command_args(cmd_args):
            return False

        # Execute conversion
        subprocess.run(
            cmd_args, check=True, capture_output=True, text=True, timeout=300  # 5 minute timeout to prevent hanging
        )
        return True
    except subprocess.TimeoutExpired:
        logging.warning("LibreOffice conversion timed out after 5 minutes")
        return False
    except Exception as e:
        logging.warning(f"LibreOffice conversion failed: {e}")
        return False


def _convert_with_docx2pdf(source_filename: str, output_pdf_filename: str) -> bool:
    if source_filename.endswith(".docx") and convert_vars.can_convert_to_pdf:
        try:
            import docx2pdf  # type: ignore

            docx2pdf.convert(source_filename, output_pdf_filename)
            logging.info(f"New file saved: {output_pdf_filename}")
            return True
        except Exception as e:
            logging.warning(f"\nConvert error: {e}")
    return False


def _handle_conversion_failure(source_filename: str) -> None:
    error_msg = (
        f"Error. A temporary file {source_filename} was created in the output folder but cannot be converted "
        f"to pdf on operating system: {platform.system()}.\n"
        "Please install LibreOffice for cross-platform PDF support."
    )
    # Check if we should suggest MS Word
    is_win_or_mac = platform.system().lower() in ["windows", "darwin"]
    libreoffice_bin = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice_bin and is_win_or_mac:
        error_msg += " This does work with MS Word installed for .docx files."
    logging.warning(error_msg)


def _cleanup_temp_file(filename: str) -> None:
    if not convert_vars.args.debug:
        try:
            os.remove(filename)
        except OSError:
            pass


def _rename_libreoffice_output(source_filename: str, output_pdf_filename: str) -> None:
    # LibreOffice outputs to the same name as source but with .pdf
    default_out = str(Path(source_filename).with_suffix(".pdf"))
    if os.path.normpath(default_out) != os.path.normpath(output_pdf_filename):
        if os.path.exists(output_pdf_filename):
            os.remove(output_pdf_filename)
        os.rename(default_out, output_pdf_filename)
    logging.info(f"New file saved: {output_pdf_filename}")


def convert_to_pdf(source_filename: str, output_pdf_filename: str) -> None:
    """Convert a document file (ODF, DOCX) to PDF using LibreOffice or docx2pdf.

    Note: The source file is preserved after conversion, as it's typically an output
    file that should be kept alongside the PDF, not a temporary file.
    """
    logging.debug(
        f" --- source_file = {source_filename} convert to {output_pdf_filename}\n--- starting pdf conversion now."
    )

    # 1. Attempt using LibreOffice
    if _convert_with_libreoffice(source_filename, output_pdf_filename):
        _rename_libreoffice_output(source_filename, output_pdf_filename)
        # Don't delete the source file - we want to keep both ODF and PDF
        return

    # 2. Fallback to docx2pdf
    if _convert_with_docx2pdf(source_filename, output_pdf_filename):
        _cleanup_temp_file(source_filename)
        return

    # 3. Everything failed
    _handle_conversion_failure(source_filename)
    # Don't delete the source file even on failure - it may still be useful


def create_edition_from_template(
    layout: str, language: str = "en", template: str = "bridge", version: str = "3.0", edition: str = "webapp"
) -> None:

    # Get the list of available translation files
    yaml_files = get_files_from_of_type(os.sep.join([convert_vars.BASE_PATH, "source"]), "yaml")
    if not yaml_files:
        return

    mapping: Dict[str, Any] = get_mapping_for_edition(yaml_files, version, language, edition, template, layout)

    if not mapping:
        logging.warning(
            f"No mapping file found for version: {version}, lang: {language}, edition: {edition},"
            f" template: {template}, layout: {layout}"
        )
        # return

    # Get the language data from the correct language file (checks vars.args.language to select the correct file)
    language_data: Dict[str, Dict[str, str]] = get_language_data(yaml_files, language, version, edition)

    # Transform the language data into the template mapping
    language_dict: Dict[str, str] = map_language_data_to_template(language_data)

    # Get meta data from language data
    meta: Dict[str, str] = get_meta_data(language_data)

    if not meta:
        logging.error("No metadata found. Cannot proceed.")
        return

    template_doc: str = get_template_for_edition(layout, template, edition)
    if template_doc == "None":
        return
    file_name, file_extension = os.path.splitext(template_doc)
    logging.debug(f"template_doc: {template_doc}")
    # Name output file with correct edition, component, language & version
    output_file: str = rename_output_file(file_extension, template, layout, meta)
    ensure_folder_exists(os.path.dirname(output_file))

    # Work with docx/odt file (and maybe convert to pdf afterwards)
    if file_extension in (".docx", ".odt"):
        language_dict.update(mapping)
        if file_extension == ".docx":
            # Get the input (template) document
            doc = get_docx_document(template_doc)
            if doc:
                doc = replace_docx_inline_text(doc, language_dict)
                doc.save(output_file)
        else:
            save_odt_file(template_doc, language_dict, output_file)

        if convert_vars.args.pdf:
            # If file type is pdf, then convert the generated file to pdf
            pdf_output_file = str(Path(output_file).with_suffix(".pdf"))
            convert_to_pdf(output_file, pdf_output_file)
    elif file_extension == ".idml":
        language_dict.update(mapping)
        save_idml_file(template_doc, language_dict, output_file)

    logging.info(f"New file saved: {output_file}")


def valid_meta(meta: Dict[str, Any], language: str, edition: str, version: str, template: str, layout: str) -> bool:
    if not has_translation_for_edition(meta, language):
        logging.warning(
            f"Translation in {language} does not exist for edition: {edition}, version: {version} "
            "or the translation choices are missing from the meta -> languages section in the mappings file"
        )
        return False

    if not has_template_for_edition(meta, template) and not convert_vars.args.inputfile:
        logging.warning(
            f"The template: {template} does not exist for edition: {edition}, version: {version} "
            "or the template choices are missing from the meta templates section in the mappings file"
        )
        return False

    if not has_layout_for_edition(meta, layout) and not convert_vars.args.inputfile:
        logging.warning(
            f"The layout: {layout} does not exist for edition: {edition}, version: {version} "
            "or the layout choices are missing from the meta -> layouts section in the mappings file"
        )
        return False
    return True


def has_translation_for_edition(meta: Dict[str, Any], language: str) -> bool:
    if meta and "languages" in meta and language in meta["languages"]:
        return True
    return False


def has_template_for_edition(meta: Dict[str, Any], template: str) -> bool:
    if meta and "templates" in meta and template in meta["templates"]:
        return True
    return False


def has_layout_for_edition(meta: Dict[str, Any], layout: str) -> bool:
    if meta and "layouts" in meta and layout in meta["layouts"]:
        return True
    return False


def ensure_folder_exists(folder_path: str) -> None:
    """Check if folder exists and if not, create folders recursively."""
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)


def main() -> None:
    convert_vars.args = parse_arguments(sys.argv[1:])
    set_logging()
    logging.debug(" --- args = %s", str(convert_vars.args))

    set_can_convert_to_pdf()
    libreoffice_available = bool(shutil.which("libreoffice") or shutil.which("soffice"))
    can_make_pdf = convert_vars.can_convert_to_pdf or libreoffice_available
    if convert_vars.args.pdf and not can_make_pdf and not convert_vars.args.debug:
        logging.error(
            "Cannot convert to pdf on this system. "
            "Pdf conversion is available on Windows and Mac (with MS Word), or on any system with LibreOffice."
        )
        return

    # Create output files
    for edition in get_valid_edition_choices():
        for layout in get_valid_layout_choices():
            for language in get_valid_language_choices():
                for template in get_valid_templates():
                    for version in get_valid_version_choices():
                        create_edition_from_template(layout, language, template, version, edition)


def parse_arguments(input_args: List[str]) -> argparse.Namespace:
    """Parse and validate the input arguments. Return object containing argument values."""
    description = "Tool to output OWASP Cornucopia playing cards into different file types and languages. "
    description += "\nExample usage: $ scripts/convert.py --pdf -lt guide -l es -v 2.2"
    description += "\nExample usage: $ scripts/convert.py -t tarot -l en -lt cards  -v 5.0 -e eop  -i "
    description += "./resources/templates/eop_ver_cards_tarot_lang.idml -o ./output/eop-5.0-cards-en.idml"
    description += "\nExample usage: c:\\cornucopia\\scripts\\convert.py -t bridge -lt cards -l fr -v 2.2 -o"
    description += " my_output_folder\\owasp_cornucopia_edition_version_layout_language_template.idml"
    parser = argparse.ArgumentParser(
        description=description, formatter_class=argparse.RawTextHelpFormatter, exit_on_error=False
    )
    parser.add_argument(
        "-i",
        "--inputfile",
        type=validate_filepath_arg,
        default="",
        help=(
            "Input (template) file to use."
            f"\nDefault={convert_vars.DEFAULT_TEMPLATE_FILENAME}.(docx|idml)"
            "\nTemplate type is dependent on the file (-o) specified."
        ),
    )
    parser.add_argument(
        "-v",
        "--version",
        type=is_valid_string_argument,
        required=False,
        default="latest",
        help=(
            "Output version to produce. [`all`, `latest`, `1.0`, `1.1`, `2.2`, `3.0`] "
            "\nFor the Website edition:"
            "\nVersion 3.0 will deliver cards mapped to ASVS 5.0"
            "\nVersion 2.2 will deliver cards mapped to ASVS 4.0"
            "\nFor the Mobile edition:"
            "\nVersion 1.0 will deliver cards mapped to MASVS 2.0"
            "\nVersion 1.1 will deliver cards mapped to MASVS 2.0"
            "\nVersion all will deliver all versions of cornucopia"
            "\nVersion latest will deliver the latest deck versions of cornucopia"
            "\nYou can also specify another version explicitly if needed. "
            "If so, there needs to be a yaml file in the source folder where the name contains "
            "the version code. Eg. edition-template-ver-lang.yaml"
        ),
    )
    parser.add_argument(
        "-o",
        "--outputfile",
        default="",
        type=validate_filepath_arg,
        help=(
            "Specify a path and name of output file to generate. (caution: existing file will be overwritten). "
            f"\nEg. {convert_vars.DEFAULT_OUTPUT_FILENAME}.(docx|pdf|idml)"
        ),
    )
    parser.add_argument(
        "-p",
        "--pdf",
        action="store_true",
        default=False,
        help=(
            "Whether to generate a pdf in addition to the printable document. "
            "Does not generate a pdf by default. Only docx can be converted to pdf for the moment."
        ),
    )
    parser.add_argument(
        "-d",
        "--debug",
        action="store_true",
        help="Output additional information to debug script",
    )
    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument(
        "-l",
        "--language",
        type=is_valid_string_argument,
        default="en",
        help=(
            "Output language to produce. [`en`, `es`, `fr`, `nl`, `no-nb`, `pt-pt`, `pt-br`, `it`, `ru`] "
            "you can also specify your own language file. If so, there needs to be a yaml "
            "file in the source folder where the name ends with the language code. Eg. edition-template-ver-lang.yaml"
        ),
    )
    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument(
        "-t",
        "--template",
        type=is_valid_string_argument,
        default="bridge",
        help=(
            "From which template to produce the document. [`bridge`, `tarot` or `tarot_qr`]\n"
            "Templates need to be added to ./resource/templates or specified with (-i or --inputfile)\n"
            "Bridge cards are 2.25 x 3.5 inch and have the mappings printed on them, \n"
            "tarot cards are 2.75 x 4.75 (71 x 121 mm) inch large, \n"
            "qr cards have a QRCode that points to an maintained list.\n"
            "You can also speficy your own template. If so, there needs to be a file in the templates folder "
            "where the name contains the template code. Eg. owasp_cornucopia_edition_ver_layout_template_lang.idml"
        ),
    )
    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument(
        "-e",
        "--edition",
        type=is_valid_string_argument,
        default="all",
        help=(
            "Output decks to produce. [`all`, `webapp` or `mobileapp`]\n"
            "The various Cornucopia decks. `web` will give you the Website App edition.\n"
            "`mobileapp` will give you the Mobile App edition.\n"
            "You can also speficy your own edition. If so, there needs to be a yaml "
            "file in the source folder where the name contains the edition code. Eg. edition-template-ver-lang.yaml"
        ),
    )

    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument(
        "-lt",
        "--layout",
        type=is_valid_string_argument,
        default="all",
        help=(
            "Document layouts to produce. [`all`, `guide`, `leaflet` or `cards`]\n"
            "The various Cornucopia document layouts.\n"
            "`cards` will output the high quality print card deck.\n"
            "`guide` will generate the docx guide with the low quality print deck.\n"
            "`leaflet` will output the high quality print leaflet.\n"
            "You can also speficy your own layout. If so, there needs to be a yaml "
            "file in the source folder where the name contains the layout code. Eg. edition-layout-ver-lang.yaml"
        ),
    )
    try:
        args = parser.parse_args(input_args)
    except argparse.ArgumentError as exc:
        # sys.tracebacklimit = 0
        logging.error(exc.message)
        sys.exit()
    return args


def is_valid_string_argument(argument: str) -> str:
    if len(argument) > 255:
        raise argparse.ArgumentTypeError("The option can not have more the 255 char.")
    if not re.match(
        r"^[\u0600-\u06FF\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF\uFDF2\uFDF3\uFDF4\uFDFD\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF\uFF66-\uFF9Fー々〆〤\u3400-\u4DBF\uF900-\uFAFF\u0900-\u097F\u0621-\u064A\u0660-\u0669\u4E00-\u9FFF\u0E00-\u0E7F«»฿ฯ๏๚๛\u0400-\u04FF\u0500-\u052F\u2DE0-\u2DFF\uA640-\uA69FЮ́ю́Я́я́\u0370-\u03FF\u1F00-\u1FFFA-Za-zÀ-ÖØ-öø-ÿĀ-ž0-9._\--ءآأؤإئابةتثجحخدذرزسشصضطظعغفقكلمنهوي ًٌٍَُِّْٰﷲﷴﷺﷻ ٠١٢٣٤٥٦٧٨٩ \s]+$",  # noqa: E501
        argument,
    ):
        raise argparse.ArgumentTypeError(
            "The option can only contain a-z letters, numbers, periods, dash or underscore"
        )
    return argument


def is_valid_argument_list(arguments: List[str]) -> Any:
    if not isinstance(arguments, List):
        return arguments
    for argument in arguments:
        is_valid_string_argument(argument)
    return arguments


def get_document_paragraphs(doc: Any) -> List[Any]:
    paragraphs = list(doc.paragraphs)
    l1 = len(paragraphs)
    for table in doc.tables:
        paragraphs += get_paragraphs_from_table_in_doc(table)
    l2 = len(paragraphs)
    if not len(paragraphs):
        logging.error("No paragraphs found in doc")
    logging.debug(f" --- count doc paragraphs = {l1}, with table paragraphs = {l2}")
    return paragraphs


def get_docx_document(docx_file: str) -> Any:
    """Open the file and return the docx document."""
    import docx  # type: ignore[import-untyped]

    if os.path.isfile(docx_file):
        return docx.Document(docx_file)
    else:
        logging.error("Could not find file at: %s", str(docx_file))
        # Create a blank document if it fails
        return docx.Document()


def get_files_from_of_type(path: str, ext: str) -> List[str]:
    """Get a list of files from a specified folder recursively, that have the specified extension."""
    files = []
    for root, dirnames, filenames in os.walk(path):
        for filename in fnmatch.filter(filenames, "*." + str(ext)):
            files.append(os.path.join(root, filename))
    if not files:
        logging.error("No language files found in folder: %s", str(os.sep.join([convert_vars.BASE_PATH, "source"])))
        return files
    logging.debug(
        "%s%s", f" --- found {len(files)} files of type {ext}. Showing first few:\n* ", str("\n* ".join(files[:3]))
    )
    return files


def get_find_replace_list(meta: Dict[str, str], template: str, layout: str) -> List[Tuple[str, str]]:
    ll: List[Tuple[str, str]] = [
        ("_edition", "_" + meta["edition"].lower()),
        ("_layout", "_" + layout.lower()),
        ("_document_template", "_" + template.lower()),
        ("_lang", "_" + meta["language"].lower()),
        ("_ver", "_" + meta["version"].lower()),
    ]
    return ll


def get_full_tag(cat_id: str, id: str, tag: str) -> str:
    if cat_id == "Common":
        full_tag = "${{{}}}".format("_".join([cat_id, id]))
    else:
        full_tag = "${{{}}}".format("_".join([cat_id, id, tag]))
    return full_tag


def get_mapping_for_edition(
    yaml_files: List[str], version: str, language: str, edition: str, template: str, layout: str
) -> Dict[str, Any]:
    mapping_data: Dict[str, Any] = get_mapping_data_for_edition(yaml_files, language, version, edition)
    if not mapping_data:
        logging.warning("No mapping file found")
        return {}
    if "meta" not in mapping_data or not valid_meta(mapping_data["meta"], language, edition, version, template, layout):
        logging.warning("Metadata is missing or invalid in mapping file")
        return {}
    try:
        mapping_data = build_template_dict(mapping_data)
    except Exception as e:
        logging.warning(f"Failed to build template mapping: {e}")
    return mapping_data


def get_mapping_data_for_edition(
    yaml_files: List[str],
    language: str,
    version: str = "3.0",
    edition: str = "webapp",
) -> Dict[Any, Dict[Any, Any]]:
    """Get the raw data of the replacement text from correct yaml file"""
    data: Dict[Any, Dict[Any, Any]] = {}
    logging.debug(
        " --- Starting get_mapping_data_for_edition() for edition: "
        f"{edition} , language: {language} and version: {version} "
        f" with mapping to version {get_valid_mapping_for_version(version, edition)}"
    )
    mappingfile: str = ""
    for file in yaml_files:
        if is_yaml_file(file) and is_mapping_file_for_version(file, version, edition):
            mappingfile = file
    if not mappingfile:
        return data

    with open(mappingfile, "r", encoding="utf-8") as f:
        try:
            data = yaml.safe_load(f)
        except yaml.YAMLError as e:
            logging.info(f"Error loading yaml file: {mappingfile}. Error = {e}")
            data = {}
    if "meta" in data.keys() and "component" in data["meta"].keys() and data["meta"]["component"] == "mappings":
        logging.debug(" --- found mappings file: " + os.path.split(mappingfile)[1])
    else:
        logging.debug(" --- found source file, but it was missing metadata: " + os.path.split(mappingfile)[1])
        if "meta" in list(data.keys()):
            meta_keys = data["meta"].keys()
            logging.debug(f" --- data.keys() = {data.keys()}, data[meta].keys() = {meta_keys}")
        data = {}
    logging.debug(f" --- Len = {len(data)}.")
    return data


def build_template_dict(input_data: Dict[str, Any]) -> Dict[str, Any]:
    """Build template dictionary from the input data"""
    data: Dict[str, Any] = {"meta": get_meta_data(input_data)}
    for key in list(k for k in input_data.keys() if k != "meta"):
        for paragraphs in input_data[key]:
            text_type = ""
            if key == "suits":
                text_type = "cards"
            if key == "paragraphs":
                text_type = "sentences"
            logging.debug(f" --- key = {key}.")
            logging.debug(f" --- suit name = {paragraphs['name']}")
            logging.debug(f" --- suit id = {is_valid_string_argument(paragraphs['id'])}")
            full_tag = "${{{}}}".format("_".join([is_valid_string_argument(paragraphs["id"]), "suit"]))
            logging.debug(f" --- suit tag = {full_tag}")
            if data["meta"]["component"] == "cards":
                data[full_tag] = paragraphs["name"]
            for paragraph in paragraphs[text_type]:
                for tag, text_output in paragraph.items():
                    logging.debug(f" --- tag = {tag}")
                    logging.debug(f" --- text = {text_output}")
                    logging.debug(f" --- paragraph = {paragraph}")
                    full_tag = get_full_tag(
                        is_valid_string_argument(paragraphs["id"]), is_valid_string_argument(paragraph["id"]), tag
                    )
                    logging.debug(f" --- full tag = {full_tag}")
                    data[full_tag] = check_make_list_into_text(text_output)
    return data


def get_meta_data(data: Dict[str, Any]) -> Dict[str, Any]:
    meta: Dict[str, Any] = {}
    if not data or "meta" not in data:
        logging.error("Could not find meta tag in the language data.")
        return meta

    raw_meta = data["meta"]
    if not isinstance(raw_meta, dict):
        logging.error("Meta tag is not a dictionary.")
        return meta

    valid_keys = ("edition", "component", "language", "version", "languages", "layouts", "templates")
    for key in valid_keys:
        if key in raw_meta:
            value = raw_meta[key]
            # Simple validation: must be string or list of strings
            if isinstance(value, str) and value.strip():
                meta[key] = value
            elif isinstance(value, list) and all(isinstance(v, str) for v in value):
                meta[key] = value

    logging.info(f" --- extracted meta data = {meta}")
    return meta


def get_paragraphs_from_table_in_doc(doc_table: Any) -> List[Any]:
    paragraphs: List[Any] = []
    for row in doc_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if len(paragraph.runs):
                    paragraphs.append(paragraph)
            for t2 in cell.tables:
                paragraphs += get_paragraphs_from_table_in_doc(t2)
    return paragraphs


def get_language_data(
    yaml_files: List[str],
    language: str,
    version: str = "3.0",
    edition: str = "webapp",
) -> Dict[Any, Any]:
    """Get the raw data of the replacement text from correct yaml file"""
    language_file: str = ""
    for file in yaml_files:
        if is_yaml_file(file) and is_lang_file_for_version(file, version, language, edition):
            language_file = file
    if not language_file:
        logging.error(f"Did not find translation for version: {version}, lang: {language}, edition: {edition}")
        return {}

    logging.debug(f" --- Loading language file: {language_file}")
    with open(language_file, "r", encoding="utf-8") as f:
        try:
            data = yaml.safe_load(f)
        except yaml.YAMLError as e:
            logging.error(f"Error loading yaml file: {language_file}. Error = {e}")
            data = {}

    if not data or "meta" not in data:
        logging.error(f"Invalid or empty language file: {language_file}")
        return {}

    return cast(Dict[Any, Any], data)


def is_mapping_file_for_version(path: str, version: str, edition: str) -> bool:
    return (
        os.path.basename(path).find("mappings") >= 0
        and os.path.basename(path).find(edition) >= 0
        and os.path.basename(path).find(version) >= 0
    )


def is_lang_file_for_version(path: str, version: str, lang: str, edition: str) -> bool:
    filename = os.path.basename(path).lower()
    # Support both -en. and -en_US. style
    lang_patterns = ["-" + lang.lower() + ".", "-" + lang.lower().replace("-", "_") + "."]
    has_lang = any(filename.find(p) != -1 for p in lang_patterns)

    return (
        filename.find(edition.lower()) != -1
        and filename.find(version.lower()) != -1
        and has_lang
        and filename.find("mappings") == -1
    )


def is_yaml_file(path: str) -> bool:
    return os.path.splitext(path)[1] in (".yaml", ".yml")


def map_language_data_to_template(input_data: Dict[str, Any]) -> Dict[str, str]:
    try:
        data = build_template_dict(input_data)
    except Exception as e:
        logging.warning(f"Could not build valid template mapping. The Yaml file is not valid. Got exception: {e}")
        data = input_data

    if convert_vars.args.debug:
        debug_txt = " --- Translation data showing First 4 (key: text):\n* "
        debug_txt += "\n* ".join(l1 + ": " + str(data[l1]) for l1 in list(data.keys())[:4])
        logging.debug(debug_txt)
        debug_txt = " --- Translation data showing Last 4 (key: text):\n* "
        debug_txt += "\n* ".join(l1 + ": " + str(data[l1]) for l1 in list(data.keys())[-4:])
        logging.debug(debug_txt)
    return data


def get_replacement_mapping_value(k: str, v: str, el_text: str) -> str:
    reg_str: str = (
        "^(OWASP MASTG|OWASP MASVS|OWASP SCP|OWASP ASVS|OWASP AppSensor|CAPEC™|SAFECODE)\u2028"
        + k.replace("$", "\\$").strip()
        + "$"
    )
    if re.match(reg_str, el_text.strip()):
        if len(v) >= 38:
            return el_text[: el_text.find("\u2028")] + ": " + v
        else:
            return el_text.replace(k, v)
    return ""


def get_replacement_value_from_dict(el_text: str, replacement_values: List[Tuple[str, str]]) -> str:
    # Fast path: if no $ and no OWASP, likely no tags
    if "$" not in el_text and "OWASP" not in el_text:
        return el_text

    for k, v in replacement_values:
        # Avoid expensive regex if key is not even in text
        if k.strip() not in el_text:
            continue

        el_new = get_replacement_mapping_value(k, v, el_text)
        if el_new:
            return el_new
        reg = r"(?<!\S)" + re.escape(k.strip()) + r"(?!\S)"
        el_text = re.sub(reg, v, el_text)
    return el_text


def get_suit_tags_and_key(key: str, edition: str) -> Tuple[List[str], str]:
    # Short tags to match the suits in the template documents
    suit_tags: List[str] = []
    suit_key: str = ""
    if key == "suits" and edition == "webapp":
        suit_tags = ["VE", "AT", "SM", "AZ", "CR", "C", "WC"]
        suit_key = "cards"
    if key == "suits" and edition == "mobileapp":
        suit_tags = ["PC", "AA", "NS", "RS", "CRM", "CM", "WC"]
        suit_key = "cards"
    elif key == "paragraphs":
        suit_tags = ["Common"]
        suit_key = "sentences"
    return suit_tags, suit_key


def get_template_for_edition(layout: str = "guide", template: str = "bridge", edition: str = "webapp") -> str:
    template_doc: str
    args_input_file: str = convert_vars.args.inputfile
    sfile_ext = "idml"
    if layout == "guide":
        sfile_ext = "odt"
    if args_input_file:
        # Input file was specified
        if os.path.isabs(args_input_file):
            template_doc = args_input_file
        elif os.path.isfile(convert_vars.BASE_PATH + os.sep + args_input_file):
            template_doc = os.path.normpath(convert_vars.BASE_PATH + os.sep + args_input_file)
        elif os.path.isfile(convert_vars.BASE_PATH + os.sep + args_input_file.replace(".." + os.sep, "")):
            template_doc = os.path.normpath(
                convert_vars.BASE_PATH + os.sep + args_input_file.replace(".." + os.sep, "")
            )
        elif args_input_file.find("..") == -1 and os.path.isfile(
            convert_vars.BASE_PATH + os.sep + ".." + os.sep + args_input_file
        ):
            template_doc = os.path.normpath(convert_vars.BASE_PATH + os.sep + ".." + os.sep + args_input_file)
        elif os.path.isfile(convert_vars.BASE_PATH + os.sep + args_input_file.replace("scripts" + os.sep, "")):
            template_doc = os.path.normpath(
                convert_vars.BASE_PATH + os.sep + args_input_file.replace("scripts" + os.sep, "")
            )
        else:
            template_doc = args_input_file
            logging.debug(f" --- Template_doc NOT found. Input File = {args_input_file}")
    else:
        # No input file specified - using defaults
        template_doc = os.path.normpath(
            convert_vars.BASE_PATH
            + os.sep
            + convert_vars.DEFAULT_TEMPLATE_FILENAME.replace("edition", edition)
            .replace("layout", layout)
            .replace("document_template", template)
            + "."
            + sfile_ext
        )

    template_doc = template_doc.replace("\\ ", " ")
    template_doc = str(Path(sanitize_filepath(template_doc)))
    if os.path.isfile(template_doc):
        template_doc = check_fix_file_extension(template_doc, sfile_ext)
        logging.debug(f" --- Returning template_doc = {template_doc}")
        return template_doc
    else:
        logging.error(f"Source file not found: {template_doc}. Please ensure file exists and try again.")
        return "None"


def get_valid_layout_choices() -> List[str]:
    layouts = []
    if convert_vars.args.layout.lower() == "all" or convert_vars.args.layout == "":
        for layout in convert_vars.LAYOUT_CHOICES:
            if layout not in ("all", "guide"):
                layouts.append(layout)
            if layout == "guide" and convert_vars.args.edition.lower() in "webapp":
                layouts.append(layout)
    else:
        layouts.append(convert_vars.args.layout)
    return layouts


def get_valid_language_choices() -> List[str]:
    languages = []
    if convert_vars.args.language.lower() == "all":
        for language in convert_vars.LANGUAGE_CHOICES:
            if language not in ("all", "template"):
                languages.append(language)
    elif convert_vars.args.language == "":
        languages.append("en")
    else:
        languages.append(convert_vars.args.language)
    return languages


def get_valid_version_choices() -> List[str]:
    versions = []
    edition: str = convert_vars.args.edition.lower()
    if convert_vars.args.version.lower() == "all":
        for version in convert_vars.VERSION_CHOICES:
            if version not in ("all", "latest") and not get_valid_mapping_for_version(version, edition) == "":
                versions.append(version)
    elif convert_vars.args.version == "" or convert_vars.args.version == "latest":
        for version in convert_vars.LATEST_VERSION_CHOICES:
            if not get_valid_mapping_for_version(version, edition) == "":
                versions.append(version)
    else:
        versions.append(convert_vars.args.version)

    if not versions:
        logging.debug(f"No deck with version: {convert_vars.args.version} for edition: {edition} exists")
    return versions


def get_valid_mapping_for_version(version: str, edition: str) -> str:
    return ConvertVars.EDITION_VERSION_MAP.get(edition, {}).get(version, "")


def get_valid_templates() -> List[str]:
    templates = []
    if convert_vars.args.template.lower() == "all":
        for template in [t for t in convert_vars.TEMPLATE_CHOICES if t not in "all"]:
            templates.append(template)
    elif convert_vars.args.template == "":
        templates.append("bridge")
        templates.append("tarot_qr")
    else:
        templates.append(convert_vars.args.template)
    return templates


def get_valid_edition_choices() -> List[str]:
    editions = []
    if convert_vars.args.edition.lower() == "all" or not convert_vars.args.edition.lower():
        for edition in convert_vars.EDITION_CHOICES:
            if edition not in "all":
                editions.append(edition)
    if convert_vars.args.edition and convert_vars.args.edition not in "all":
        editions.append(convert_vars.args.edition)
    return editions


def group_number_ranges(data: List[str]) -> List[str]:
    if len(data) < 2 or len([s for s in data if not str(s).isnumeric()]):
        return data
    list_ranges: List[str] = []
    data_numbers = [int(s) for s in data]
    for k, g in groupby(enumerate(data_numbers), lambda x: x[0] - x[1]):
        group: List[int] = list(map(itemgetter(1), g))
        group = list(map(int, group))
        if group[0] == group[-1]:
            list_ranges.append(str(group[0]))
        else:
            list_ranges.append(str(group[0]) + "-" + str(group[-1]))
    return list_ranges


def save_docx_file(doc: Any, output_file: str) -> None:
    ensure_folder_exists(os.path.dirname(output_file))
    doc.save(output_file)


def save_odt_file(template_doc: str, language_dict: Dict[str, str], output_file: str) -> None:
    # Get the output path and temp output path to put the temp xml files
    output_path = os.path.join(convert_vars.BASE_PATH, "output")
    temp_output_path = os.path.join(output_path, "temp_odt")
    # Ensure the output folder and temp output folder exist
    ensure_folder_exists(temp_output_path)
    logging.debug(" --- temp_folder for extraction of xml files = %s", str(temp_output_path))

    # Unzip source xml files and place in temp output folder
    with zipfile.ZipFile(template_doc) as odt_archive:
        _safe_extractall(odt_archive, temp_output_path)

    # ODT text is usually in content.xml and sometimes styles.xml
    targets = ["content.xml", "styles.xml"]
    replacement_values = sort_keys_longest_to_shortest(language_dict)

    for target in targets:
        xml_file = os.path.join(temp_output_path, target)
        if os.path.exists(xml_file):
            replace_text_in_xml_file(xml_file, replacement_values)

    # Zip the files as an odt file in output folder
    logging.debug(" --- finished replacing text in xml files. Now zipping into odt file")
    zip_dir(temp_output_path, output_file)

    # If not debugging, delete temp folder and files
    if not convert_vars.args.debug and os.path.exists(temp_output_path):
        shutil.rmtree(temp_output_path, ignore_errors=True)


def save_idml_file(template_doc: str, language_dict: Dict[str, str], output_file: str) -> None:
    # Get the output path and temp output path to put the temp xml files
    output_path = convert_vars.BASE_PATH + os.sep + "output"
    temp_output_path = output_path + os.sep + "temp"
    # Ensure the output folder and temp output folder exist
    ensure_folder_exists(temp_output_path)
    logging.debug(" --- temp_folder for extraction of xml files = %s", str(temp_output_path))

    # Unzip source xml files and place in temp output folder
    with zipfile.ZipFile(template_doc) as idml_archive:
        _safe_extractall(idml_archive, temp_output_path)
        logging.debug(" --- namelist of first few files in archive = %s", str(idml_archive.namelist()[:5]))

    xml_files = get_files_from_of_type(temp_output_path, "xml")
    replacement_values = sort_keys_longest_to_shortest(language_dict)
    # Only Stories files have content to update
    for file in fnmatch.filter(xml_files, "*Stories*Story*"):
        if os.path.getsize(file) == 0:
            continue
        replace_text_in_xml_file(file, replacement_values)

    # Zip the files as an idml file in output folder
    logging.debug(" --- finished replacing text in xml files. Now zipping into idml file")
    zip_dir(temp_output_path, output_file)

    # If not debugging, delete temp folder and files
    if not convert_vars.args.debug and os.path.exists(temp_output_path):
        shutil.rmtree(temp_output_path, ignore_errors=True)


def set_can_convert_to_pdf() -> bool:
    operating_system: str = sys.platform.lower()
    can_convert = operating_system.find("win") != -1 or operating_system.find("darwin") != -1
    convert_vars.can_convert_to_pdf = can_convert
    logging.debug(f" --- operating system = {operating_system}, can_convert_to_pdf = {convert_vars.can_convert_to_pdf}")
    return can_convert


def set_logging() -> None:
    logging.basicConfig(
        format="%(asctime)s %(filename)s | %(levelname)s | %(funcName)s | %(message)s",
    )
    if convert_vars.args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
    else:
        logging.getLogger().setLevel(logging.INFO)


def sort_keys_longest_to_shortest(replacement_dict: Dict[str, str]) -> List[Tuple[str, str]]:
    new_list = list((k, v) for k, v in replacement_dict.items())
    return sorted(new_list, key=lambda s: len(s[0]), reverse=True)


def remove_short_keys(replacement_dict: Dict[str, str], min_length: int = 8) -> Dict[str, str]:
    data2: Dict[str, str] = {}
    for key, value in replacement_dict.items():
        if len(key) >= min_length:
            data2[key] = value
    logging.debug(
        " --- Making template. Removed card_numbers. len replacement_dict = "
        f"{len(replacement_dict)}, len data2 = {len(data2)}"
    )
    return data2


def rename_output_file(file_extension: str, template: str, layout: str, meta: Dict[str, str]) -> str:
    """Rename output file replacing place-holders from meta dict (edition, component, language, version)."""
    args_output_file: str = convert_vars.args.outputfile
    logging.debug(f" --- args_output_file = {args_output_file}")
    if args_output_file:
        # Output file is specified as an argument
        if os.path.isabs(args_output_file):
            output_filename = args_output_file
        else:
            output_filename = str(Path(convert_vars.BASE_PATH + os.sep + args_output_file))
    else:

        # No output file specified - using default
        output_filename = str(
            Path(convert_vars.BASE_PATH + os.sep + convert_vars.DEFAULT_OUTPUT_FILENAME + file_extension)
        )

    logging.debug(f" --- output_filename before fix extension = {output_filename}")
    output_filename = check_fix_file_extension(output_filename, file_extension)
    logging.debug(f" --- output_filename AFTER fix extension = {output_filename}")

    # Do the replacement of filename place-holders with metadata
    find_replace = get_find_replace_list(meta, template, layout)
    f = os.path.basename(output_filename)
    for r in find_replace:
        f = f.replace(*r)
    output_filename = os.path.dirname(output_filename) + os.sep + f
    output_filename = str(Path(sanitize_filepath(output_filename)))

    logging.debug(f" --- output_filename = {output_filename}")
    return output_filename


def replace_docx_inline_text(doc: Any, data: Dict[str, str]) -> Any:
    """Replace the text in the docx document."""
    logging.debug(" --- starting docx_replace")
    replacement_values = list(data.items())
    paragraphs = get_document_paragraphs(doc)
    for p in paragraphs:
        runs_text = "".join(r.text for r in p.runs)
        if runs_text.strip() == "":
            continue
        for key, val in replacement_values:
            replaced_key = False
            for i, run in enumerate(p.runs):
                if run.text.find(key) != -1:
                    p.runs[i].text = run.text.replace(key, val)
                    replaced_key = True
                    runs_text = runs_text.replace(key, val)
            if not replaced_key:
                if runs_text.find(key) != -1:
                    runs_text = runs_text.replace(key, val)
                    for i, r in enumerate(p.runs):
                        p.runs[i].text = ""
                    p.runs[0].text = runs_text

    logging.debug(" --- finished replacing text in doc")
    return doc


def _find_xml_elements(tree: Any) -> List[ElTree.Element]:
    """Identify elements likely to contain text to replace for IDML and ODT."""
    namespaces = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
        "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
    }
    elements = tree.findall(".//Content")
    elements.extend(tree.findall(".//text:p", namespaces))
    elements.extend(tree.findall(".//text:span", namespaces))
    if not elements:
        return cast(List[ElTree.Element], tree.findall(".//*"))
    return cast(List[ElTree.Element], elements)


def replace_text_in_xml_file(filename: str, replacement_values: List[Tuple[str, str]]) -> None:
    logging.debug(f" --- starting xml_replace for {filename}")
    try:
        tree = DefusedElTree.parse(filename)
    except Exception as e:
        logging.error(f"Failed to parse XML file {filename}: {e}")
        return

    root = tree.getroot()
    if root is None:
        logging.error(f" --- The XML file has no root element: {filename}")
        return

    elements_to_check = _find_xml_elements(tree)

    modified = False
    for el in elements_to_check:
        if el.text:
            new_text = get_replacement_value_from_dict(el.text, replacement_values)
            if new_text != el.text:
                el.text = new_text
                modified = True

    if modified:
        try:
            with open(filename, "bw") as f:
                f.write(ElTree.tostring(root, encoding="utf-8"))
        except Exception as e:
            logging.error(f"Failed to save modified XML file {filename}: {e}")


def zip_dir(path: str, zip_filename: str) -> None:
    """Zip all the files recursively from path into zip_filename (excluding root path)"""
    with zipfile.ZipFile(zip_filename, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for root, _, files in os.walk(os.path.normpath(path)):
            for file in files:
                f = str(Path(os.path.join(root, file)))
                zip_file.write(f, f[len(path) :])  # noqa: E203


if __name__ == "__main__":
    convert_vars: ConvertVars = ConvertVars()
    main()
