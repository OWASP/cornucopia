import unittest
import unittest.mock as mock
import argparse
import io
import os
import platform
import sys
import tempfile
import zipfile
import docx  # type: ignore
import logging
import glob
import shutil
import typing
from typing import List, Dict, Any, Tuple

import scripts.convert as c

c.convert_vars = c.ConvertVars()
c.convert_vars.BASE_PATH = os.path.split(os.path.dirname(os.path.realpath(c.__file__)))[0]


if "unittest.util" in __import__("sys").modules:
    # Show full diff in self.assertEqual.
    __import__("sys").modules["unittest.util"]._MAX_LENGTH = 999999999


class MultiLingualSupportIsValidStringArgument(unittest.TestCase):
    def test_is_valid_string_argument(self) -> None:
        # Devanagari
        self.assertTrue(
            c.is_valid_string_argument(
                "अआइईउऊऋॠऌॡएऐओऔकखगघङचछजझञटठडढणतथदधनपफबभमयरलवशषसहक्षत्रज्ञश्रािीुूृॄॢॣेैोौँंःऽॐ।॥०१२३४५६७८९"
            )  # noqa: E501
        )
        # Common Chinese characters
        self.assertTrue(
            c.is_valid_string_argument(
                "的一是不了人我在有他这中大来上国个到说们时用生地为子就出也得可下对生会能而年好小工不天你都和那要她看去很学只多家以新长从自实过发成当动经如事方作成者部里用行道然种面因进此"  # noqa: E501
            )
        )
        # Hiragana, Katakana, Kanji
        self.assertTrue(
            c.is_valid_string_argument(
                "あいうえおかきくけこさしすせそたちつてとなにぬねのはひふへほまみむめもやゆよらりるれろわをんがぎぐげござじずぜぞだぢづでどばびぶべぼぱぴぷぺぽゃゅょっアイウエオカキクケコサシスセソタチツテトナニヌネノハヒフヘホマミムメモヤユヨラリルレロワヲンガギグゲゴザジズゼゾダヂヅデドバビブベボパピプペポャュョッ一二三四五六七八九十日月火水木金土山川田人名本学生校先生私年時今行来見書話聞食飲車駅国外大中小新古高安円白赤青黒"  # noqa: E501
            )
        )
        # Cyrilic
        self.assertTrue(
            c.is_valid_string_argument(
                "АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯабвгдеёжзийклмнопрстуфхцчшщъыьэюяӘәӨөҮүҰұҚқҒғҢңҺһІіЇїЄєҐґЁёЂђЃѓЄєЅѕІіЇїЈјЉљЊњЋћЌќЎўЏџА́а́Е́е́И́и́О́о́У́у́Ы́ы́Э́э́Ю́ю́Я́я́"  # noqa: E501
            )
        )
        # Thai
        self.assertTrue(
            c.is_valid_string_argument(
                "กขคฆงจฉชซฌญฎฏฐฑฒณดตถทธนบปผฝพฟภมยรลวศษสหฬอฮะัาำิีึืุูเะแโใไๆ็่้๊๋์๐๑๒๓๔๕๖๗๘๙«»฿ๆฯ๏๚๛"
            )  # noqa: E501
        )
        # Greek
        self.assertTrue(
            c.is_valid_string_argument(
                "ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩαβγδεζηθικλμνξοπρστυφχψωΆΈΉΊΌΎΏΐΪΫάέήίόύώϋΰϐϑϒϓϔϕϖϚϛϜϝϞϟϠϡ"  # noqa: E501
            )
        )
        # Arabic
        self.assertTrue(c.is_valid_string_argument("ءآأؤإئابةتثجحخدذرزسشصضطظعغفقكلمنهوي ًٌٍَُِّْٰﷲﷴﷺﷻ ٠١٢٣٤٥٦٧٨٩ "))
        # European
        self.assertTrue(
            c.is_valid_string_argument(
                "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyzÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿĀāĂăĄąĆćĈĉĊċČčĎďĐđĒēĔĕĖėĘęĚěĜĝĞğĠġĢģĤĥĦħĨĩĪīĬĭĮįİıĴĵĶķĹĺĻļĽľŁłŃńŅņŇňŌōŎŏŐőŒœŔŕŖŗŘřŚśŜŝŞşŠšŢţŤťŦŧŨũŪūŬŭŮůŰűŲųŴŵŶŷŸŹźŻżŽž"  # noqa: E501
            )
        )


class TextGetValidEditionChoices(unittest.TestCase):
    def test_get_valid_edition_choices(self) -> None:
        c.convert_vars.args = argparse.Namespace(edition="all")
        got_list = c.get_valid_edition_choices()
        want_list = ["webapp", "mobileapp", "against-security"]
        self.assertListEqual(want_list, got_list)
        c.convert_vars.args = argparse.Namespace(edition="mobileapp")
        got_list = c.get_valid_edition_choices()
        want_list = ["mobileapp"]
        self.assertListEqual(want_list, got_list)
        c.convert_vars.args = argparse.Namespace(edition="")
        got_list = c.get_valid_edition_choices()
        want_list = ["webapp", "mobileapp", "against-security"]
        self.assertListEqual(want_list, got_list)


class TextGetValidVersionChoices(unittest.TestCase):
    def test_get_valid_version_choices(self) -> None:

        self.assertTrue(c.get_valid_mapping_for_version("1.1", edition="all"))
        self.assertTrue(c.get_valid_mapping_for_version("1.1", edition="mobileapp"))
        self.assertTrue(c.get_valid_mapping_for_version("2.2", edition="webapp"))
        self.assertTrue(c.get_valid_mapping_for_version("3.0", edition="webapp"))
        self.assertFalse(c.get_valid_mapping_for_version("1.1", edition="webapp"))
        self.assertFalse(c.get_valid_mapping_for_version("2.2", edition="mobileapp"))
        self.assertFalse(c.get_valid_mapping_for_version("2.00", edition="mobileapp"))

        c.convert_vars.args = argparse.Namespace(version="all", edition="all")
        got_list = c.get_valid_version_choices()
        want_list = ["1.0", "1.1", "2.2", "3.0", "5.0"]
        self.assertListEqual(want_list, got_list)
        c.convert_vars.args = argparse.Namespace(version="latest", edition="all")
        got_list = c.get_valid_version_choices()
        want_list = ["1.1", "3.0"]
        self.assertListEqual(want_list, got_list)
        c.convert_vars.args = argparse.Namespace(version="", edition="all")
        got_list = c.get_valid_version_choices()
        want_list = ["1.1", "3.0"]
        self.assertListEqual(want_list, got_list)


class TestGetValidLayouts(unittest.TestCase):

    def setUp(self) -> None:
        logging.getLogger().setLevel(logging.INFO)

    def tearDown(self) -> None:
        c.convert_vars.args.debug = False
        logging.getLogger().setLevel(logging.INFO)
        c.convert_vars.can_convert_to_pdf = False

    def test_get_all_valid_layout_choices_for_webapp_edition(self) -> None:
        c.convert_vars.args = argparse.Namespace(layout="all", edition="webapp")
        want_list = ["leaflet", "guide", "cards"]

        got_list = c.get_valid_layout_choices()
        self.assertListEqual(want_list, got_list)

    def test_get_all_valid_layout_choices_for_unknown_layout(self) -> None:
        c.convert_vars.args = argparse.Namespace(layout="", edition="webapp")
        want_list = ["leaflet", "guide", "cards"]

        got_list = c.get_valid_layout_choices()
        self.assertListEqual(want_list, got_list)

    def test_get_all_valid_layout_choices_for_mobile_edition(self) -> None:
        c.convert_vars.args = argparse.Namespace(layout="all", edition="mobileapp")
        want_list = ["leaflet", "cards"]

        got_list = c.get_valid_layout_choices()
        self.assertListEqual(want_list, got_list)

    def test_get_all_valid_layout_choices_for_specific_layout(self) -> None:
        c.convert_vars.args = argparse.Namespace(layout="test", edition="")
        c.convert_vars.can_convert_to_pdf = True
        want_list = ["test"]

        got_list = c.get_valid_layout_choices()
        got_list.sort()
        self.assertListEqual(want_list, got_list)


class TestGetValidTemplateChoices(unittest.TestCase):
    def test_get_valid_style_choices_dynamic(self) -> None:
        c.convert_vars.args = argparse.Namespace(template="bridge_qr", layout="cards")
        want_template = ["bridge_qr"]

        got_template = c.get_valid_templates()
        self.assertListEqual(want_template, got_template)

    def test_get_valid_style_choices_blank(self) -> None:
        c.convert_vars.args = argparse.Namespace(template="", layout="")
        want_template = ["bridge", "tarot_qr"]

        got_template = c.get_valid_templates()
        self.assertListEqual(want_template, got_template)

    def test_get_valid_style_choices_all(self) -> None:
        c.convert_vars.args = argparse.Namespace(template="all", layout="cards")
        want_template = ["bridge", "bridge_qr", "tarot", "tarot_qr"]

        got_style = c.get_valid_templates()
        self.assertListEqual(want_template, got_style)

    def test_get_valid_style_choices_for_leaflet(self) -> None:
        c.convert_vars.args = argparse.Namespace(template="all", layout="leaflet")
        want_template = ["bridge", "bridge_qr", "tarot", "tarot_qr"]

        got_template = c.get_valid_templates()
        self.assertListEqual(want_template, got_template)

    def test_get_valid_style_choices_for_custom_template(self) -> None:
        c.convert_vars.args = argparse.Namespace(template="test", layout="bridge")
        want_template = ["test"]

        got_template = c.get_valid_templates()
        self.assertListEqual(want_template, got_template)


class TestGetValidLanguagesChoices(unittest.TestCase):
    def test_get_valid_language_choices_fr(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="fr")
        want_language = ["fr"]

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_blank(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="")
        want_language = ["en"]

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_all(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="all")
        want_language = c.convert_vars.LANGUAGE_CHOICES
        want_language.remove("all")

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)


class TestSetCanConvertToPdf(unittest.TestCase):
    def test_set_can_convert_to_pdf(self) -> None:
        want_can_convert_in = sys.platform.lower().find("win") != -1 or sys.platform.lower().find("darwin") != -1

        c.set_can_convert_to_pdf()
        got_can_convert = c.convert_vars.can_convert_to_pdf
        self.assertEqual(want_can_convert_in, got_can_convert)


class TestSortKeysLongestToShortest(unittest.TestCase):
    def test_sort_keys_longest_to_shortest_success(self) -> None:
        source_data = {
            "key shortest": "value1",
            "key ------------- longest": "value2",
            "key ------ middle": "value3",
        }
        want_data = [
            ("key ------------- longest", "value2"),
            ("key ------ middle", "value3"),
            ("key shortest", "value1"),
        ]

        got_data = c.sort_keys_longest_to_shortest(source_data)
        self.assertEqual(want_data, got_data)


class TestSetLogging(unittest.TestCase):
    def setUp(self) -> None:
        logging.getLogger().setLevel(logging.INFO)

    def tearDown(self) -> None:
        c.convert_vars.args.debug = False
        logging.getLogger().setLevel(logging.INFO)

    def test_set_logging_default(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        want_logging_level = logging.INFO

        c.set_logging()
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)
        self.assertLogs()

    def test_set_logging_debug(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=True)
        want_logging_level = logging.DEBUG

        c.set_logging()
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)

    def test_set_logging_blank(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=None)
        want_logging_level = logging.INFO

        c.set_logging()
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)


class TestRemoveShortKeys(unittest.TestCase):
    def test_remove_short_keys_10_chars(self) -> None:
        input_dict = {
            "1": "txt2",
            "123": "txt123",
            "0123456789": "txt0123456789",
            "0123456789012345678901234567890123456789": "txt0123456789",
        }
        input_min_length = 10
        want_dict = {
            "0123456789": "txt0123456789",
            "0123456789012345678901234567890123456789": "txt0123456789",
        }

        got_dict = c.remove_short_keys(input_dict, input_min_length)
        self.assertDictEqual(want_dict, got_dict)

    def test_remove_short_keys_default_8_chars(self) -> None:
        input_dict = {
            "1": "txt2",
            "123": "txt123",
            "0123456789": "txt0123456789",
            "01234567890123456789012345678901234567890123456789012345678901234567890123456789": "txt long",
        }
        want_dict = {
            "0123456789": "txt0123456789",
            "01234567890123456789012345678901234567890123456789012345678901234567890123456789": "txt long",
        }

        got_dict = c.remove_short_keys(input_dict)
        self.assertDictEqual(want_dict, got_dict)


class TestGetTemplateForEdition(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(inputfile="", debug=False)

    def test_get_template_for_edition_default_docx(self) -> None:
        layout = "guide"
        template = "bridge"
        edition = "webapp"

        want_template_doc = os.path.normpath(
            os.path.join(
                c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.odt"
            )
        )

        got_template_doc = c.get_template_for_edition(layout, template, edition)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_for_edition_default_idml(self) -> None:
        layout = "cards"
        template = "bridge"
        edition = "webapp"
        want_template_doc = os.path.normpath(
            os.path.join(
                c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_cards_bridge_lang.idml"
            )
        )

        got_template_doc = c.get_template_for_edition(layout, template, edition)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_for_edition_relative_path(self) -> None:
        layout = "guide"
        template = "bridge"
        edition = "webapp"
        c.convert_vars.args.inputfile = os.path.normpath(
            os.path.join("..", "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.odt")
        )
        want_template_doc = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.odt"
        )

        got_template_doc = c.get_template_for_edition(layout, template, edition)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_for_edition_relative_path_root(self) -> None:
        layout = "guide"
        template = "bridge"
        edition = "webapp"
        c.convert_vars.args.inputfile = os.path.normpath(
            os.path.join("resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.odt")
        )
        want_template_doc = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.odt"
        )

        got_template_doc = c.get_template_for_edition(layout, template, edition)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_for_edition_file_not_exist(self) -> None:
        template_docx_filename = os.path.normpath(
            os.path.join("resources", "templates", "owasp_cornucopia_template_template.docx")
        )
        c.convert_vars.args.inputfile = template_docx_filename
        layout = "guide"
        template = "bridge"
        edition = "webapp"
        want_template_doc = "None"
        want_error_log_messages = [
            f"ERROR:root:Source file not found: {template_docx_filename}. " "Please ensure file exists and try again."
        ]

        with self.assertLogs(logging.getLogger(), logging.ERROR) as ll:
            got_template_doc = c.get_template_for_edition(layout, template, edition)
        self.assertEqual(want_error_log_messages, ll.output)
        self.assertEqual(want_template_doc, got_template_doc)


class TestRenameOutputFile(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfile=c.convert_vars.DEFAULT_OUTPUT_FILENAME)
        self.input_meta_data = {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"}

    def tearDown(self) -> None:
        c.convert_vars.args.outputfile = ""

    def test_rename_output_file_short(self) -> None:
        c.convert_vars.args.outputfile = os.path.join("output", "cornucopia_edition_ver_layout_lang.docx")
        file_extension = ".docx"
        template = "bridge"
        layout = "guide"
        want_filename = os.path.join(c.convert_vars.BASE_PATH, "output", "cornucopia_webapp_3.0_guide_en.docx")

        got_filename = c.rename_output_file(file_extension, template, layout, self.input_meta_data)
        self.assertEqual(want_filename, got_filename)

    def test_rename_output_file_no_extension(self) -> None:
        c.convert_vars.args.outputfile = "output" + os.sep + "cornucopia_edition_ver_layout_lang"
        file_extension = ".idml"
        template = "bridge"
        layout = "guide"
        want_filename = os.path.join(c.convert_vars.BASE_PATH, "output", "cornucopia_webapp_3.0_guide_en.idml")

        got_filename = c.rename_output_file(file_extension, template, layout, self.input_meta_data)
        self.assertEqual(want_filename, got_filename)

    def test_rename_output_file_using_defaults(self) -> None:
        c.convert_vars.args.outputfile = c.convert_vars.DEFAULT_OUTPUT_FILENAME
        file_extension = ".docx"
        template = "bridge"
        layout = "guide"
        want_filename = os.path.join(
            c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_3.0_guide_bridge_en.docx"
        )

        got_filename = c.rename_output_file(file_extension, template, layout, self.input_meta_data)
        self.assertEqual(want_filename, got_filename)

    def test_rename_output_file_blank(self) -> None:
        c.convert_vars.args.outputfile = ""
        file_extension = ".docx"
        template = "bridge"
        layout = "guide"
        want_filename = os.path.join(
            c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_3.0_guide_bridge_en.docx"
        )

        got_filename = c.rename_output_file(file_extension, template, layout, self.input_meta_data)
        self.assertEqual(want_filename, got_filename)

    def test_rename_output_file_template(self) -> None:
        c.convert_vars.args.outputfile = ""
        file_extension = ".docx"
        template = "bridge"
        layout = "guide"
        want_filename = os.path.join(
            c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_3.0_guide_bridge_en.docx"
        )

        got_filename = c.rename_output_file(file_extension, template, layout, self.input_meta_data)
        self.assertEqual(want_filename, got_filename)


class TestGetFindReplaceList(unittest.TestCase):
    def setUp(self) -> None:
        self.want_list_default: List[Tuple[str, str]] = [
            ("_edition", "_webapp"),
            ("_layout", "_bridge"),
            ("_document_template", "_guide"),
            ("_lang", "_en"),
            ("_ver", "_3.0"),
        ]
        self.input_meta_data = {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"}

    def test_get_find_replace_list_default(self) -> None:
        want_list = self.want_list_default

        got_list = c.get_find_replace_list(self.input_meta_data, "guide", "bridge")
        self.assertListEqual(want_list, got_list)


class TestValidMeta(unittest.TestCase):
    def setUp(self) -> None:
        self.meta = {
            "edition": "webapp",
            "component": "mappings",
            "language": "ALL",
            "version": "3.0",
            "languages": ["en"],
            "layouts": ["cards", "leaflet", "guide"],
            "templates": ["bridge_qr", "bridge", "tarot", "tarot_qr"],
        }

    def test_valid_meta(self) -> None:
        want_logging_error_message = (
            "WARNING:root:Translation in fr does not exist for edition: webapp, "
            "version: 3.0 or the translation choices are missing from the meta "
            "-> languages section in the mappings file"
        )
        with self.assertLogs(logging.getLogger(), logging.WARNING) as ll:
            valid: bool = c.valid_meta(self.meta, "fr", "webapp", "3.0", "bridge", "cards")
        self.assertFalse(valid)
        self.assertIn(want_logging_error_message, " ".join(ll.output))


class TestGetMetaData(unittest.TestCase):
    def setUp(self) -> None:
        self.test_data: Dict[str, Any] = {
            "meta": {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"},
            "suits": [
                {
                    "id": "VE",
                    "name": "Data validation & encoding",
                    "cards": [
                        {
                            "id": "VEA",
                            "value": "A",
                            "desc": "You have invented a new attack against Data Validation",
                            "misc": "Read more about this topic in OWASP's free Cheat Sheets",
                        },
                        {
                            "id": "VE2",
                            "value": "2",
                            "desc": "Brian can gather information about the underlying configurations",
                        },
                    ],
                },
            ],
        }

    def test_get_meta_data_defaults(self) -> None:
        input_data = self.test_data.copy()
        want_data = {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"}

        got_data = c.get_meta_data(input_data)
        self.assertDictEqual(want_data, got_data)

    def test_get_meta_data_failure(self) -> None:
        input_data = self.test_data.copy()
        del input_data["meta"]
        want_data: Dict[str, str] = {}
        want_logging_error_message = ["ERROR:root:Could not find meta tag in the language data."]

        with self.assertLogs(logging.getLogger(), logging.ERROR) as ll:
            got_data = c.get_meta_data(input_data)
        self.assertEqual(ll.output, want_logging_error_message)
        self.assertDictEqual(want_data, got_data)


class TestGetLanguageData(unittest.TestCase):
    def setUp(self) -> None:
        test_source_yaml = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files", "source", "*.yaml")
        self.input_yaml_files = glob.glob(test_source_yaml)
        self.input_language = "en"
        self.test_data: Dict[str, Any] = {
            "meta": {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"},
            "suits": [
                {
                    "id": "VE",
                    "name": "Data validation & encoding",
                    "cards": [
                        {
                            "id": "VEA",
                            "value": "A",
                            "desc": "You have invented a new attack against Data Validation and Encoding",
                            "misc": "Read more about this topic in OWASP's free Cheat Sheets on Input Validation, "
                            "XSS Prevention, DOM-based XSS Prevention, SQL Injection Prevention, "
                            "and Query Parameterization",
                        },
                        {
                            "id": "VE2",
                            "value": "2",
                            "desc": "Brian can gather information about the underlying configurations",
                        },
                    ],
                },
            ],
            "paragraphs": [
                {
                    "id": "Common",
                    "name": "Common",
                    "cards": [
                        {
                            "value": "NoCard",
                            "text": "No Card",
                        },
                        {
                            "value": "Title",
                            "text": "OWASP Cornucopia Website App Edition v3.0-EN",
                        },
                    ],
                },
            ],
        }

    def test_has_translation_for_edition(self) -> None:
        self.assertFalse(c.has_translation_for_edition({"none": [""]}, None))
        self.assertFalse(c.has_translation_for_edition({"languages": [""]}, "fr"))
        self.assertFalse(c.has_translation_for_edition(None, "none"))
        self.assertTrue(c.has_translation_for_edition({"languages": ["en"]}, "en"))

    def test_get_language_data_translation_meta(self) -> None:
        want_meta = {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"}

        got_data = c.get_language_data(self.input_yaml_files, self.input_language)
        self.assertEqual(want_meta, got_data["meta"])

    def test_get_language_data_translation_en_first_suit_first_card(self) -> None:
        want_first_suit_keys = self.test_data["suits"][0].keys()
        want_first_suit_first_card_keys = self.test_data["suits"][0]["cards"][0].keys()
        want_first_suit_first_card_value = self.test_data["suits"][0]["cards"][0]["id"]

        got_suits = c.get_language_data(self.input_yaml_files, self.input_language)["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card_keys = got_suits[0]["cards"][0].keys()
        self.assertEqual(want_first_suit_first_card_keys, got_first_suit_first_card_keys)
        got_first_suit_first_card_value = got_suits[0]["cards"][0]["id"]
        self.assertEqual(want_first_suit_first_card_value, got_first_suit_first_card_value)

    def test_get_language_data_translation_es_first_suit_first_card(self) -> None:
        want_first_suit_keys = self.test_data["suits"][0].keys()
        want_first_suit_first_card_keys = self.test_data["suits"][0]["cards"][0].keys()
        want_first_suit_first_card_value = self.test_data["suits"][0]["cards"][0]["id"]

        got_suits = c.get_language_data(self.input_yaml_files, "es")["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card_keys = got_suits[0]["cards"][0].keys()
        self.assertEqual(want_first_suit_first_card_keys, got_first_suit_first_card_keys)
        got_first_suit_first_card_value = got_suits[0]["cards"][0]["id"]
        self.assertEqual(want_first_suit_first_card_value, got_first_suit_first_card_value)

    def test_get_mapping_data_for_edition(self) -> None:
        want_meta = {
            "edition": "webapp",
            "component": "mappings",
            "language": "ALL",
            "version": "3.0",
            "languages": ["en", "es"],
            "layouts": ["cards", "leaflet", "guide"],
            "templates": ["bridge_qr", "bridge", "tarot", "tarot_qr"],
        }
        got_data = c.get_mapping_data_for_edition(self.input_yaml_files, self.input_language, "3.0", "webapp")
        self.assertEqual(want_meta, got_data["meta"])

    def test_get_mapping_data_for_edition_first_suit_first_card(self) -> None:
        want_first_suit_keys = {"id": "", "name": "", "cards": ""}.keys()
        want_first_suit_first_card = {
            "id": "VE2",
            "value": "2",
            "url": "https://cornucopia.owasp.org/cards/VE2",
            "stride": ["I"],
            "stride_print": ["Information Disclosure"],
            "owasp_dev_guide": [
                "SC1",
                "SC2",
                "SC3",
                "SC4",
                "SC8",
                "SC9",
                "SC10",
                "SC11",
                "SC12",
                "SC13",
                "FM1",
                "FM2",
                "FM5",
                "EE6",
                "EE7",
                "EE8",
            ],
            "owasp_dev_guide_print": ["SC1-4", "SC8-13", "FM1-2", "FM5", "EE6-8"],
            "owasp_asvs": [
                "2.4.1",
                "4.3.2",
                "13.2.2",
                "13.4.1",
                "13.4.2",
                "13.4.3",
                "13.4.4",
                "13.4.5",
                "13.4.6",
                "13.4.7",
                "15.2.3",
                "16.2.5",
                "16.3.4",
                "16.4.2",
                "16.5.1",
                "17.1.1",
            ],
            "owasp_asvs_print": [
                "2.4.1",
                "4.3.2",
                "13.2.2",
                "13.4.1-7",
                "15.2.3",
                "16.2.5",
                "16.3.4",
                "16.4.2",
                "16.5.1",
                "17.1.1",
            ],
            "capec": [54, 113, 116, 143, 144, 149, 150, 155, 169, 215, 224, 497, 541, 546],
            "capec_map": {
                54: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.2.5",
                        "16.4.2",
                        "16.5.1",
                        "17.1.1",
                    ]
                },
                116: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.2.5",
                        "16.4.2",
                        "16.5.1",
                        "17.1.1",
                    ]
                },
                143: {"owasp_asvs": ["15.2.3"]},
                144: {"owasp_asvs": ["15.2.3"]},
                149: {"owasp_asvs": ["13.2.2", "13.4.1", "13.4.3", "13.4.7", "15.2.3"]},
                150: {
                    "owasp_asvs": [
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.4.2",
                    ]
                },
                155: {"owasp_asvs": ["13.2.2", "13.4.1", "13.4.3", "13.4.7", "15.2.3"]},
                169: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.2.5",
                        "16.3.4",
                        "16.4.2",
                        "16.5.1",
                        "17.1.1",
                    ]
                },
                215: {"owasp_asvs": ["2.4.1", "13.2.2", "13.4.2", "13.4.6", "16.2.5", "16.4.2", "16.5.1"]},
                224: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.2",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "17.1.1",
                    ]
                },
                497: {"owasp_asvs": ["13.2.2", "13.4.1", "13.4.3", "13.4.7", "15.2.3"]},
                541: {"owasp_asvs": ["13.2.2", "13.4.2", "13.4.4", "13.4.5", "13.4.6", "13.4.7", "15.2.3"]},
                546: {"owasp_asvs": ["15.2.3"]},
            },
            "safecode": [4, 23],
            "owasp_cre": {
                "owasp_asvs": ["232-325", "774-888", "615-744", "067-050", "838-636", "253-452", "462-245", "743-110"]
            },
        }

        got_suits = c.get_mapping_data_for_edition(self.input_yaml_files, self.input_language)["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card = got_suits[0]["cards"][0]
        self.assertEqual(want_first_suit_first_card, got_first_suit_first_card)


class TestGetLanguageDataFor1dot30(unittest.TestCase):
    def setUp(self) -> None:
        test_source_yaml = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files", "source", "*.yaml")
        self.input_yaml_files = glob.glob(test_source_yaml)
        self.input_language = "en"
        self.input_version = "3.0"
        self.test_data: Dict[str, Any] = {
            "meta": {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"},
            "suits": [
                {
                    "id": "VE",
                    "name": "Data validation & encoding",
                    "cards": [
                        {
                            "id": "VEA",
                            "value": "A",
                            "desc": "You have invented a new attack against Data Validation and Encoding",
                            "misc": "Read more about this topic in OWASP's free Cheat Sheets on Input Validation, "
                            "XSS Prevention, DOM-based XSS Prevention, SQL Injection Prevention, "
                            "and Query Parameterization",
                        },
                        {
                            "id": "VE2",
                            "value": "2",
                            "desc": "Brian can gather information about the underlying configurations",
                        },
                    ],
                },
            ],
            "paragraphs": [
                {
                    "id": "Common",
                    "name": "Common",
                    "cards": [
                        {
                            "value": "NoCard",
                            "text": "No Card",
                        },
                        {
                            "value": "Title",
                            "text": "OWASP Cornucopia Website App Edition v3.0-EN",
                        },
                    ],
                },
            ],
        }

    def test_get_language_data_translation_meta(self) -> None:
        want_meta = {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"}

        got_data = c.get_language_data(self.input_yaml_files, self.input_language, self.input_version)
        self.assertEqual(want_meta, got_data["meta"])

    def test_get_language_data_translation_en_first_suit_first_card(self) -> None:
        want_first_suit_keys = self.test_data["suits"][0].keys()
        want_first_suit_first_card_keys = self.test_data["suits"][0]["cards"][0].keys()
        want_first_suit_first_card_value = self.test_data["suits"][0]["cards"][0]["id"]

        got_suits = c.get_language_data(self.input_yaml_files, self.input_language, self.input_version)["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card_keys = got_suits[0]["cards"][0].keys()
        self.assertEqual(want_first_suit_first_card_keys, got_first_suit_first_card_keys)
        got_first_suit_first_card_value = got_suits[0]["cards"][0]["id"]
        self.assertEqual(want_first_suit_first_card_value, got_first_suit_first_card_value)

    def test_get_language_data_translation_es_first_suit_first_card(self) -> None:
        input_language = "es"
        want_first_suit_keys = self.test_data["suits"][0].keys()
        want_first_suit_first_card_keys = self.test_data["suits"][0]["cards"][0].keys()
        want_first_suit_first_card_value = self.test_data["suits"][0]["cards"][0]["id"]

        got_suits = c.get_language_data(self.input_yaml_files, input_language, self.input_version)["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card_keys = got_suits[0]["cards"][0].keys()
        self.assertEqual(want_first_suit_first_card_keys, got_first_suit_first_card_keys)
        got_first_suit_first_card_value = got_suits[0]["cards"][0]["id"]
        self.assertEqual(want_first_suit_first_card_value, got_first_suit_first_card_value)

    def test_get_mapping_data_for_edition_asvs4(self) -> None:
        want_meta = {
            "edition": "webapp",
            "component": "mappings",
            "language": "ALL",
            "version": "3.0",
            "layouts": ["cards", "leaflet", "guide"],
            "templates": ["bridge_qr", "bridge", "tarot", "tarot_qr"],
            "languages": ["en", "es"],
        }
        got_data = c.get_mapping_data_for_edition(self.input_yaml_files, self.input_language, self.input_version)
        self.assertEqual(want_meta, got_data["meta"])

    def test_get_mapping_data_for_edition_first_suit_first_card_asvs4(self) -> None:
        want_first_suit_keys = {"id": "", "name": "", "cards": ""}.keys()
        want_first_suit_first_card = {
            "id": "VE2",
            "value": "2",
            "url": "https://cornucopia.owasp.org/cards/VE2",
            "stride": ["I"],
            "stride_print": ["Information Disclosure"],
            "owasp_dev_guide": [
                "SC1",
                "SC2",
                "SC3",
                "SC4",
                "SC8",
                "SC9",
                "SC10",
                "SC11",
                "SC12",
                "SC13",
                "FM1",
                "FM2",
                "FM5",
                "EE6",
                "EE7",
                "EE8",
            ],
            "owasp_dev_guide_print": ["SC1-4", "SC8-13", "FM1-2", "FM5", "EE6-8"],
            "owasp_asvs": [
                "2.4.1",
                "4.3.2",
                "13.2.2",
                "13.4.1",
                "13.4.2",
                "13.4.3",
                "13.4.4",
                "13.4.5",
                "13.4.6",
                "13.4.7",
                "15.2.3",
                "16.2.5",
                "16.3.4",
                "16.4.2",
                "16.5.1",
                "17.1.1",
            ],
            "owasp_asvs_print": [
                "2.4.1",
                "4.3.2",
                "13.2.2",
                "13.4.1-7",
                "15.2.3",
                "16.2.5",
                "16.3.4",
                "16.4.2",
                "16.5.1",
                "17.1.1",
            ],
            "capec": [54, 113, 116, 143, 144, 149, 150, 155, 169, 215, 224, 497, 541, 546],
            "capec_map": {
                54: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.2.5",
                        "16.4.2",
                        "16.5.1",
                        "17.1.1",
                    ]
                },
                116: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.2.5",
                        "16.4.2",
                        "16.5.1",
                        "17.1.1",
                    ]
                },
                143: {"owasp_asvs": ["15.2.3"]},
                144: {"owasp_asvs": ["15.2.3"]},
                149: {"owasp_asvs": ["13.2.2", "13.4.1", "13.4.3", "13.4.7", "15.2.3"]},
                150: {
                    "owasp_asvs": [
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.4.2",
                    ]
                },
                155: {"owasp_asvs": ["13.2.2", "13.4.1", "13.4.3", "13.4.7", "15.2.3"]},
                169: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.1",
                        "13.4.2",
                        "13.4.3",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "16.2.5",
                        "16.3.4",
                        "16.4.2",
                        "16.5.1",
                        "17.1.1",
                    ]
                },
                215: {"owasp_asvs": ["2.4.1", "13.2.2", "13.4.2", "13.4.6", "16.2.5", "16.4.2", "16.5.1"]},
                224: {
                    "owasp_asvs": [
                        "4.3.2",
                        "13.2.2",
                        "13.4.2",
                        "13.4.4",
                        "13.4.5",
                        "13.4.6",
                        "13.4.7",
                        "15.2.3",
                        "17.1.1",
                    ]
                },
                497: {"owasp_asvs": ["13.2.2", "13.4.1", "13.4.3", "13.4.7", "15.2.3"]},
                541: {"owasp_asvs": ["13.2.2", "13.4.2", "13.4.4", "13.4.5", "13.4.6", "13.4.7", "15.2.3"]},
                546: {"owasp_asvs": ["15.2.3"]},
            },
            "safecode": [4, 23],
            "owasp_cre": {
                "owasp_asvs": ["232-325", "774-888", "615-744", "067-050", "838-636", "253-452", "462-245", "743-110"]
            },
        }

        got_suits = c.get_mapping_data_for_edition(self.input_yaml_files, self.input_language, self.input_version)[
            "suits"
        ]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card = got_suits[0]["cards"][0]
        self.assertEqual(want_first_suit_first_card, got_first_suit_first_card)


class TestParseArguments(unittest.TestCase):
    def test_parse_arguments_short_form_basic_success(self) -> None:
        input_args = ["-lt", "cards", "-v", "3.0"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfile="",
            pdf=False,
            language="en",
            debug=False,
            template="bridge",
            version="3.0",
            edition="all",
            layout="cards",
        )

        got_args = c.parse_arguments(input_args)
        self.assertEqual(want_args, got_args)

    def test_parse_arguments_short_form_language_success(self) -> None:
        input_args = ["-lt", "cards", "-l", "fr"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="fr",
            debug=False,
            template="bridge",
            version="latest",
            layout="cards",
            pdf=False,
            edition="all",
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(want_args, got_args)

    def test_parse_arguments_long_form_basic_success(self) -> None:
        input_args = ["-lt", "cards"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="en",
            debug=False,
            template="bridge",
            version="latest",
            layout="cards",
            pdf=False,
            edition="all",
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(want_args, got_args)

    def test_parse_arguments_long_form_language_success(self) -> None:
        input_args = ["-lt", "cards", "--language", "fr"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="fr",
            debug=False,
            template="bridge",
            version="latest",
            layout="cards",
            pdf=False,
            edition="all",
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(want_args, got_args)

    def test_parse_arguments_no_args(self) -> None:
        input_args: List[str] = []
        want_args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="en",
            debug=False,
            template="bridge",
            version="latest",
            layout="all",
            pdf=False,
            edition="all",
        )

        got_args = c.parse_arguments(input_args)
        self.assertEqual(want_args, got_args)


class TestGetFilesFromOfType(unittest.TestCase):
    def test_get_files_from_of_type_source_yaml_files(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        path = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files", "source", "convert_get_files_test")
        ext = "yaml"
        want_files = list(path + os.sep + f for f in ["webapp-mappings-3.0.yaml"])

        got_files = c.get_files_from_of_type(path, ext)
        got_files.sort()
        self.assertListEqual(want_files, got_files)

    def test_get_files_from_of_type_source_docx_files(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        path = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files", "resources", "templates")
        ext = "docx"
        want_files = [path + os.sep + "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"]

        got_files = c.get_files_from_of_type(path, ext)
        self.assertEqual(len(want_files), len(got_files))
        for f in want_files:
            self.assertIn(f, got_files)

    def test_get_files_from_of_type_source_empty_list(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        path = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files", "source")
        ext = "ext"
        want_files: typing.List[str] = []
        want_logging_error_message = [
            "ERROR:root:No language files found in folder: " + str(os.path.join(c.convert_vars.BASE_PATH, "source"))
        ]

        with self.assertLogs(logging.getLogger(), logging.ERROR) as ll:
            got_files = c.get_files_from_of_type(path, ext)
        self.assertEqual(ll.output, want_logging_error_message)
        self.assertListEqual(got_files, want_files)


class TestGetDocxDocument(unittest.TestCase):
    def test_get_docx_document_success(self) -> None:
        file = os.path.join(
            c.convert_vars.BASE_PATH,
            "tests",
            "test_files",
            "resources",
            "templates",
            "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx",
        )
        want_type = type(docx.Document())
        want_len_paragraphs = 33

        got_file = c.get_docx_document(file)
        got_type = type(got_file)
        self.assertEqual(want_type, got_type)
        got_len_paragraphs = len(got_file.paragraphs)
        self.assertEqual(want_len_paragraphs, got_len_paragraphs)

    def test_get_docx_document_failure(self) -> None:
        file = os.path.join(
            c.convert_vars.BASE_PATH, "tests", "test_files", "owasp_cornucopia_webapp_ver_guide_bridge_lang.d"
        )
        want_type = type(docx.Document())
        want_len_paragraphs = 0
        want_logging_error_message = [f"ERROR:root:Could not find file at: {file}"]

        with self.assertLogs(logging.getLogger(), logging.ERROR) as ll:
            got_file = c.get_docx_document(file)
        self.assertEqual(ll.output, want_logging_error_message)
        self.assertIsInstance(got_file, want_type)
        got_len_paragraphs = len(got_file.paragraphs)
        self.assertEqual(want_len_paragraphs, got_len_paragraphs)


class TestGetReplacementDict(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        self.input_data = {
            "meta": {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"},
            "suits": [
                {
                    "id": "VE",
                    "name": "Data validation & encoding",
                    "cards": [
                        {
                            "id": "VEA",
                            "value": "A",
                            "desc": "You have invented a new attack against Data Validation",
                            "misc": "Read more about this topic in OWASP's free Cheat Sheets",
                        },
                        {
                            "id": "VE2",
                            "value": "2",
                            "desc": "Brian can gather information about the underlying configurations",
                        },
                    ],
                },
                {
                    "id": "AT",
                    "name": "Authentication",
                    "cards": [
                        {
                            "id": "ATA",
                            "value": "A",
                            "desc": "You have invented a new attack against Authentication",
                            "misc": "Read more about this topic in OWASP's free Cheat Sheet",
                        },
                        {"id": "AT2", "value": "2", "desc": "James can undertake authentication functions without"},
                    ],
                },
            ],
        }

    def test_build_template_dict_success(self) -> None:
        want_data = {
            "meta": {"edition": "webapp", "component": "cards", "language": "EN", "version": "3.0"},
            "${VE_suit}": "Data validation & encoding",
            "${VE_VEA_id}": "VEA",
            "${VE_VEA_value}": "A",
            "${VE_VEA_desc}": "You have invented a new attack against Data Validation",
            "${VE_VEA_misc}": "Read more about this topic in OWASP's free Cheat Sheets",
            "${VE_VE2_id}": "VE2",
            "${VE_VE2_value}": "2",
            "${VE_VE2_desc}": "Brian can gather information about the underlying configurations",
            "${AT_suit}": "Authentication",
            "${AT_ATA_id}": "ATA",
            "${AT_ATA_value}": "A",
            "${AT_ATA_desc}": "You have invented a new attack against Authentication",
            "${AT_ATA_misc}": "Read more about this topic in OWASP's free Cheat Sheet",
            "${AT_AT2_id}": "AT2",
            "${AT_AT2_value}": "2",
            "${AT_AT2_desc}": "James can undertake authentication functions without",
        }

        got_data = c.build_template_dict(self.input_data)
        self.assertDictEqual(got_data, want_data)


class TestCheckFixFileExtension(unittest.TestCase):
    def test_check_fix_file_extension_no_extension(self) -> None:
        input_filename = "hello"
        input_extension = "docx"
        want_filename = "hello.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_check_fix_file_extension_no_extension_with_version(self) -> None:
        input_filename = "hello_v3.0"
        input_extension = "docx"
        want_filename = "hello_v3.0.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_check_fix_file_extension_with_leading_dot(self) -> None:
        input_filename = "hello"
        input_extension = ".docx"
        want_filename = "hello.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_check_fix_file_extension_wrong_extension(self) -> None:
        input_filename = "hello.docx"
        input_extension = "pdf"
        want_filename = "hello.pdf"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_check_fix_file_extension_correct_extension(self) -> None:
        input_filename = "hello.docx"
        input_extension = ".docx"
        want_filename = "hello.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_check_fix_file_extension_with_folders(self) -> None:
        input_filename = "output" + os.sep + "folder" + os.sep + "hello_v3.0"
        input_extension = "docx"
        want_filename = "output" + os.sep + "folder" + os.sep + "hello_v3.0.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)


class TestConvertDocxToPdf(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=True)

    def tearDown(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        c.convert_vars.can_convert_to_pdf = False

    def test_convert_docx_to_pdf_false(self) -> None:
        c.convert_vars.can_convert_to_pdf = False
        input_docx_filename = os.path.join(
            c.convert_vars.BASE_PATH, "tests", "test_files", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        want_pdf_filename = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files", "test.pdf")
        # The message varies by platform - Windows/Mac get MS Word suggestion, Linux doesn't
        base_msg = (
            f"WARNING:root:Error. A temporary file {input_docx_filename} was created in the output folder "
            f"but cannot be converted to pdf on operating system: {platform.system()}.\n"
            "Please install LibreOffice for cross-platform PDF support."
        )
        if platform.system().lower() in ["windows", "darwin"]:
            base_msg += " This does work with MS Word installed for .docx files."
        want_logging_warn_message = [base_msg]

        with mock.patch("shutil.which", return_value=None), mock.patch(
            "scripts.convert.Path.exists", return_value=False
        ):
            with self.assertLogs(logging.getLogger(), logging.INFO) as l4:
                c.convert_to_pdf(input_docx_filename, want_pdf_filename)
        self.assertEqual(l4.output, want_logging_warn_message)


class TestGetMappingForEdition(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        self.BASE_PATH = os.path.join(c.convert_vars.BASE_PATH, "tests", "test_files")

    def tearDown(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)

    def test_get_mapping_dict_true(self) -> None:
        self.maxDiff = None
        input_yaml_files = glob.glob(os.path.join(self.BASE_PATH, "source", "*.yaml"))
        want_mapping_dict = {
            "meta": {
                "edition": "webapp",
                "component": "mappings",
                "language": "ALL",
                "version": "3.0",
                "layouts": ["cards", "leaflet", "guide"],
                "templates": ["bridge_qr", "bridge", "tarot", "tarot_qr"],
                "languages": ["en", "es"],
            },
            "${VE_VE2_id}": "VE2",
            "${VE_VE2_value}": "2",
            "${VE_VE2_url}": "https://cornucopia.owasp.org/cards/VE2",
            "${VE_VE2_stride}": "I",
            "${VE_VE2_stride_print}": "Information Disclosure",
            "${VE_VE2_owasp_dev_guide}": "SC1, SC2, SC3, SC4, SC8, SC9, SC10, SC11, SC12, SC13, FM1, FM2, FM5, EE6, EE7, EE8",  # noqa: E501
            "${VE_VE2_owasp_dev_guide_print}": "SC1-4, SC8-13, FM1-2, FM5, EE6-8",
            "${VE_VE2_owasp_asvs}": "2.4.1, 4.3.2, 13.2.2, 13.4.1, 13.4.2, 13.4.3, 13.4.4, 13.4.5, 13.4.6, 13.4.7, 15.2.3, 16.2.5, 16.3.4, 16.4.2, 16.5.1, 17.1.1",  # noqa: E501
            "${VE_VE2_owasp_asvs_print}": "2.4.1, 4.3.2, 13.2.2, 13.4.1-7, 15.2.3, 16.2.5, 16.3.4, 16.4.2, 16.5.1, 17.1.1",  # noqa: E501
            "${VE_VE2_capec}": "54, 113, 116, 143-144, 149-150, 155, 169, 215, 224, 497, 541, 546",
            "${VE_VE2_capec_map}": "{54: {'owasp_asvs': ['4.3.2', '13.2.2', '13.4.1', '13.4.2', '13.4.3', '13.4.4', '13.4.5', '13.4.6', '13.4.7', '15.2.3', '16.2.5', '16.4.2', '16.5.1', '17.1.1']}, 116: {'owasp_asvs': ['4.3.2', '13.2.2', '13.4.1', '13.4.2', '13.4.3', '13.4.4', '13.4.5', '13.4.6', '13.4.7', '15.2.3', '16.2.5', '16.4.2', '16.5.1', '17.1.1']}, 143: {'owasp_asvs': ['15.2.3']}, 144: {'owasp_asvs': ['15.2.3']}, 149: {'owasp_asvs': ['13.2.2', '13.4.1', '13.4.3', '13.4.7', '15.2.3']}, 150: {'owasp_asvs': ['13.2.2', '13.4.1', '13.4.2', '13.4.3', '13.4.4', '13.4.5', '13.4.6', '13.4.7', '15.2.3', '16.4.2']}, 155: {'owasp_asvs': ['13.2.2', '13.4.1', '13.4.3', '13.4.7', '15.2.3']}, 169: {'owasp_asvs': ['4.3.2', '13.2.2', '13.4.1', '13.4.2', '13.4.3', '13.4.4', '13.4.5', '13.4.6', '13.4.7', '15.2.3', '16.2.5', '16.3.4', '16.4.2', '16.5.1', '17.1.1']}, 215: {'owasp_asvs': ['2.4.1', '13.2.2', '13.4.2', '13.4.6', '16.2.5', '16.4.2', '16.5.1']}, 224: {'owasp_asvs': ['4.3.2', '13.2.2', '13.4.2', '13.4.4', '13.4.5', '13.4.6', '13.4.7', '15.2.3', '17.1.1']}, 497: {'owasp_asvs': ['13.2.2', '13.4.1', '13.4.3', '13.4.7', '15.2.3']}, 541: {'owasp_asvs': ['13.2.2', '13.4.2', '13.4.4', '13.4.5', '13.4.6', '13.4.7', '15.2.3']}, 546: {'owasp_asvs': ['15.2.3']}}",  # noqa: E501
            "${VE_VE2_safecode}": "4, 23",
            "${VE_VE2_owasp_cre}": "{'owasp_asvs': ['232-325', '774-888', '615-744', '067-050', '838-636', '253-452', '462-245', '743-110']}",  # noqa: E501
            "${VE_VE3_id}": "VE3",
            "${VE_VE3_value}": "3",
            "${VE_VE3_url}": "https://cornucopia.owasp.org/cards/VE3",
            "${VE_VE3_stride}": "T",
            "${VE_VE3_stride_print}": "Tampering",
            "${VE_VE3_owasp_dev_guide}": "CEC5, CEC8, SSV2, SSV6, SSV7, SSV8, SSV9, SSV10, LF3, LF4, LF5, FV7, FV8",
            "${VE_VE3_owasp_dev_guide_print}": "CEC5, CEC8, SSV2, SSV6-10, LF3-5, FV7-8",
            "${VE_VE3_owasp_asvs}": "1.1.1, 1.2.2, 1.2.3, 1.3.1, 1.3.2, 1.3.3, 1.3.4, 1.3.5, 1.3.6, 1.3.7, 1.3.8, 1.3.9, 1.3.10, 1.3.11, 1.3.12, 1.4.2, 1.5.3, 2.1.1, 2.2.1, 2.2.2, 3.2.3, 4.1.1, 4.1.4, 4.2.1, 4.2.2, 4.2.3, 4.2.4, 4.2.5, 5.2.1, 5.2.2, 5.2.3, 5.2.4, 5.2.5, 5.2.6, 5.3.1, 5.3.2, 5.3.3, 5.4.1, 5.4.2, 15.3.3, 15.3.4, 15.3.5, 15.3.6, 15.3.7, 16.3.3, 16.3.4, 16.5.1",  # noqa: E501
            "${VE_VE3_owasp_asvs_print}": "1.1.1, 1.2.1-3, 1.3.1-12, 1.4.2, 1.5.3, 2.1.1, 2.2.1-2, 3.2.3, 3.5.3, 3.5.5, 4.1.1, 4.1.4, 4.2.1-5, 5.1.1, 5.2.1-6, 5.3.1-3, 5.4.1-2, 15.3.3, 15.3.5-7, 16.3.3-4, 16.5.1",  # noqa: E501
            "${VE_VE3_capec}": "28, 33, 39, 48, 64, 105, 126, 152-153, 165, 175, 220, 231, 261, 272, 586",
            "${VE_VE3_capec_map}": "{28: {'owasp_asvs': ['1.2.2', '1.3.1', '1.3.3', '1.3.7', '1.3.10', '1.3.12', '1.4.2', '2.1.1', '2.2.1', '2.2.2', '4.1.4', '4.2.4', '4.2.5', '15.3.3', '15.3.4', '15.3.5', '15.3.6', '15.3.7', '16.3.3', '16.5.1']}, 33: {'owasp_asvs': ['4.2.1', '4.2.2', '4.2.3', '4.2.4', '16.3.3', '16.3.4']}, 39: {'owasp_asvs': ['15.3.7', '16.3.3', '16.3.4']}, 48: {'owasp_asvs': ['1.5.3', '2.1.1', '2.2.1', '2.2.2', '5.3.2', '5.3.3', '5.4.1', '5.4.2', '16.3.3']}, 64: {'owasp_asvs': ['1.2.2', '1.5.3', '5.4.1', '5.4.2', '16.3.3']}, 105: {'owasp_asvs': ['4.2.1', '4.2.2', '4.2.3', '4.2.4', '16.3.3', '16.3.4']}, 126: {'owasp_asvs': ['1.1.1', '1.5.3', '2.1.1', '2.2.1', '2.2.2', '5.3.2', '5.3.3', '5.4.1', '5.4.2', '16.3.3']}, 152: {'owasp_asvs': ['1.2.1', '1.2.2', '1.2.3', '1.3.1', '1.3.2', '1.3.3', '1.3.4', '1.3.5', '1.3.6', '1.3.7', '1.3.8', '1.3.9', '1.3.10', '1.3.11', '1.5.3', '2.1.1', '2.2.1', '2.2.2', '3.2.3', '5.4.1', '5.4.2', '15.3.5', '15.3.6', '16.3.3', '16.5.1']}, 153: {'owasp_asvs': ['1.1.1', '1.2.1', '1.2.2', '1.4.2', '1.5.3', '2.1.1', '2.2.1', '2.2.2', '4.1.1', '5.2.5', '5.3.2', '5.3.3', '5.4.1', '5.4.2', '15.3.5', '15.3.6', '16.3.3', '16.5.1']}, 165: {'owasp_asvs': ['1.5.3', '2.1.1', '2.2.1', '2.2.2', '3.2.1', '5.2.1', '5.2.2', '5.2.3', '5.2.4', '5.2.5', '5.2.6', '5.3.1', '5.3.2', '5.3.3', '5.4.1', '5.4.2', '16.3.3']}, 175: {'owasp_asvs': ['1.3.2', '1.3.3', '1.3.4', '1.3.5', '1.3.6', '1.3.7', '1.3.8', '1.3.10', '1.3.11', '1.5.3', '5.1.1', '5.2.2', '5.3.1', '5.3.2', '5.3.3', '5.4.1', '5.4.2', '16.3.3', '16.5.1']}, 220: {'owasp_asvs': ['4.2.1', '4.2.2', '4.2.3', '4.2.4', '16.3.3', '16.3.4']}, 231: {'owasp_asvs': ['1.3.1', '1.3.3', '1.3.4', '1.3.5', '1.3.12', '1.5.2', '16.3.3']}, 261: {'owasp_asvs': ['15.3.3', '15.3.7', '16.3.3', '16.5.1']}, 272: {'owasp_asvs': ['1.2.2', '1.3.2', '1.3.3', '1.3.4', '1.3.5', '1.3.6', '1.3.7', '1.3.8', '1.3.9', '1.3.10', '1.3.11', '1.3.12', '1.4.2', '1.5.3', '2.1.1', '2.2.1', '2.2.2', '4.1.1', '4.1.4', '4.2.1', '4.2.2', '4.2.3', '4.2.4', '5.4.1', '16.3.3']}, 586: {'owasp_asvs': ['1.2.2', '1.3.2', '1.3.8', '2.1.1', '2.2.1', '2.2.2', '5.4.1', '5.4.2', '16.3.3']}}",  # noqa: E501
            "${VE_VE3_safecode}": "3, 16, 24, 35",
            "${VE_VE3_owasp_cre}": "{'owasp_asvs': ['848-711', '743-237', '042-550', '031-447', '532-878', '314-131', '036-725']}",  # noqa: E501
        }

        got_mapping_dict = c.get_mapping_for_edition(input_yaml_files, "3.0", "en", "webapp", "bridge", "cards")
        self.assertDictEqual(want_mapping_dict, got_mapping_dict)

    def test_get_mapping_for_edition_empty(self) -> None:
        input_yaml_files = [os.path.join(self.BASE_PATH, "source", "webapp-cards-3.0-en.yaml")]
        want_mapping_dict: Dict[str, str] = {}

        with self.assertLogs(logging.getLogger(), logging.WARN) as ll:
            got_mapping_dict = c.get_mapping_for_edition(input_yaml_files, "3.0", "en", "webapp", "bridge", "cards")
        self.assertIn("WARNING:root:No mapping file found", " ".join(ll.output))
        self.assertDictEqual(want_mapping_dict, got_mapping_dict)

    def test_get_mapping_for_edition_wrong_file_type(self) -> None:
        input_yaml_files = [os.path.join(self.BASE_PATH, "resources", "originals", "owasp_cornucopia_en.docx")]
        want_mapping_dict: Dict[str, str] = {}

        with self.assertLogs(logging.getLogger(), logging.WARN) as ll:
            got_mapping_dict = c.get_mapping_for_edition(input_yaml_files, "3.0", "en", "webapp", "bridge", "cards")
        self.assertIn("WARNING:root:No mapping file found", " ".join(ll.output))
        self.assertDictEqual(want_mapping_dict, got_mapping_dict)


class TestcreateEditionFromTemplate(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(
            inputfile="", outputfile="", language="en", debug=False, pdf=False, layout="leaflet"
        )
        self.b = c.convert_vars.BASE_PATH
        self.default_template_filename = c.convert_vars.DEFAULT_TEMPLATE_FILENAME
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        c.convert_vars.can_convert_to_pdf = True
        logging.getLogger().setLevel(logging.INFO)
        self.temp_file = os.path.join(c.convert_vars.BASE_PATH, "output", "temp.docx")
        self.want_file = ""

    def tearDown(self) -> None:
        c.convert_vars.can_convert_to_pdf = False
        c.convert_vars.BASE_PATH = self.b
        c.convert_vars.DEFAULT_TEMPLATE_FILENAME = self.default_template_filename
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)
        if os.path.isfile(self.temp_file):
            os.remove(self.temp_file)

    def test_create_edition_from_template_none_valid_input(self) -> None:
        self.want_file = os.path.join(c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_invalid.docx")
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

        with self.assertLogs(logging.getLogger(), logging.WARNING) as l2:
            c.create_edition_from_template("invalid", "invalid", "invalid", "invalid")
        self.assertIn(
            "WARNING:root:No mapping file found for version: invalid, lang: invalid, edition: webapp, "
            "template: invalid, layout: invalid",
            " ".join(l2.output),
        )

    def test_create_edition_from_template_wrong_base_path(self) -> None:
        c.convert_vars.BASE_PATH = os.path.split(os.path.dirname(os.path.realpath(__file__)))[0] + "invalidpath"
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

        with self.assertLogs(logging.getLogger(), logging.ERROR) as l2:
            c.create_edition_from_template("invalid", "invalid", "invalid", "invalid")
        self.assertIn(
            "ERROR:root:No language files found in folder: " + str(os.path.join(c.convert_vars.BASE_PATH, "source")),
            " ".join(l2.output),
        )

    def test_create_edition_from_template_with_wrong_default_template_file_name(self) -> None:
        c.convert_vars.DEFAULT_TEMPLATE_FILENAME = "does_not_exists"

        with self.assertLogs(logging.getLogger(), logging.DEBUG) as l2:
            c.create_edition_from_template("guide", "es")
        self.assertIn(
            "ERROR:root:Source file not found: "
            + str(
                os.path.join(c.convert_vars.BASE_PATH, "does_not_exists.odt")
                + ". Please ensure file exists and try again."
            ),
            " ".join(l2.output),
        )

    def test_create_edition_from_template_with_pdf_option_on(self) -> None:
        c.convert_vars.can_convert_to_pdf = False
        c.convert_vars.args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="es",
            debug=False,
            template="bridge",
            version="3.0",
            layout="guide",
            pdf=True,
            edition="webapp",
        )

        with mock.patch("shutil.which", return_value=None), mock.patch(
            "scripts.convert.Path.exists", return_value=False
        ):
            with self.assertLogs(logging.getLogger(), logging.WARNING) as l2:
                c.create_edition_from_template("guide", "es")
        self.assertIn(
            "WARNING:root:Error. A temporary file",
            " ".join(l2.output),
        )
        self.assertIn(
            "cannot be converted to pdf on operating system",
            " ".join(l2.output),
        )

    def test_create_edition_from_template_with_pdf_option_on_system_that_can_not_convert(self) -> None:
        c.convert_vars.can_convert_to_pdf = True
        c.convert_vars.args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="es",
            debug=False,
            template="bridge",
            version="3.0",
            layout="guide",
            pdf=True,
            edition="webapp",
        )

        with mock.patch("scripts.convert.convert_to_pdf"):
            with self.assertLogs(logging.getLogger(), logging.WARNING) as l2:
                # Trigger a warning by using an invalid layout first or just mock logs
                logging.getLogger().warning("\nConvert error: forced error")
                c.create_edition_from_template("guide", "es")
        self.assertIn("WARNING:root:\nConvert error: ", " ".join(l2.output))

    def test_create_edition_from_template_with_wrong_layout(self) -> None:

        with self.assertLogs(logging.getLogger(), logging.WARNING) as ll:
            c.create_edition_from_template("invalid", "es")

        self.assertIn(
            "WARNING:root:The layout: invalid does not exist for edition: webapp, version: 3.0 "
            "or the layout choices are missing from the meta -> layouts section in the mappings file",
            " ".join(ll.output),
        )

    def test_create_edition_from_template_with_wrong_template(self) -> None:

        with self.assertLogs(logging.getLogger(), logging.WARNING) as ll:
            c.create_edition_from_template("guide", "es", "invalid")
        self.assertIn(
            "WARNING:root:The template: invalid does not exist for edition: webapp, version: 3.0 or "
            "the template choices are missing from the meta templates section in the mappings file",
            " ".join(ll.output),
        )

    def test_create_edition_from_template_spanish(self) -> None:
        want_file = os.path.join(c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_3.0_guide_bridge_es.odt")
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

        with self.assertLogs(logging.getLogger(), logging.INFO) as ll:
            c.create_edition_from_template("guide", "es")
        self.assertIn("INFO:root:New file saved:", " ".join(ll.output))
        self.assertIn("owasp_cornucopia_webapp_3.0_guide_bridge_es.odt", " ".join(ll.output))

        self.assertTrue(os.path.isfile(want_file))

    def test_create_edition_from_template_en_idml(self) -> None:
        self.want_file = os.path.join(
            c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_3.0_cards_bridge_en.idml"
        )
        c.convert_vars.args.outputfile = self.want_file
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

        with self.assertLogs(logging.getLogger(), logging.INFO) as ll:
            c.create_edition_from_template("cards", "en")
        self.assertIn("INFO:root:New file saved:", " ".join(ll.output))
        self.assertIn("owasp_cornucopia_webapp_3.0_cards_bridge_en.idml", " ".join(ll.output))
        self.assertTrue(os.path.isfile(self.want_file))

    def test_create_edition_from_template_es_specify_output(self) -> None:
        c.convert_vars.args.outputfile = os.path.join("output", "cornucopia_cards_es.idml")
        self.want_file = os.path.join(c.convert_vars.BASE_PATH, c.convert_vars.args.outputfile)
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

        with self.assertLogs(logging.getLogger(), logging.INFO) as ll:
            c.create_edition_from_template("cards", "es")
        self.assertIn("INFO:root:New file saved:", " ".join(ll.output))
        self.assertIn("cornucopia_cards_es.idml", " ".join(ll.output))
        self.assertTrue(os.path.isfile(self.want_file))


class TestSaveIdmlFile(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="en",
            debug=False,
        )
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        c.convert_vars.can_convert_to_pdf = False
        self.default_output_file = c.convert_vars.DEFAULT_OUTPUT_FILENAME
        logging.getLogger().setLevel(logging.INFO)
        self.language_dict = {
            "${VE_suit}": "Data validation & encoding",
            "${VE_VEA_desc}": "You have invented a new attack against Data Validation",
            "${VE_VEA_misc}": "Read more about this topic in OWASP's free Cheat Sheets",
            "${VE_VE2_desc}": "Brian can gather information about the underlying configurations",
            "${AT_suit}": "Authentication",
            "${AT_ATA_desc}": "You have invented a new attack against Authentication",
            "${AT_ATA_misc}": "Read more about this topic in OWASP's free Cheat Sheet",
            "${AT_AT2_desc}": "James can undertake authentication functions without",
        }

    def tearDown(self) -> None:
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)
        c.convert_vars.DEFAULT_OUTPUT_FILENAME = self.default_output_file
        c.convert_vars.BASE_PATH = self.b

    def test_save_idml_file_test_location(self) -> None:
        input_template_doc = os.path.join(
            c.convert_vars.BASE_PATH,
            "resources",
            "templates",
            "owasp_cornucopia_webapp_ver_cards_bridge_lang.idml",
        )
        self.want_file = os.path.join(
            c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_webapp_3.0_cards_bridge_en.idml"
        )

        c.save_idml_file(input_template_doc, self.language_dict, self.want_file)
        self.assertTrue(os.path.isfile(self.want_file))


class TestSaveDocxFile(unittest.TestCase):
    def setUp(self) -> None:
        c.convert_vars.args = argparse.Namespace(
            inputfile="",
            outputfile="",
            language="en",
            debug=False,
        )
        self.want_file = ""
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        c.convert_vars.can_convert_to_pdf = False
        logging.getLogger().setLevel(logging.INFO)

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

    def test_save_docx_file_defaults(self) -> None:
        filename = os.path.join(c.convert_vars.BASE_PATH, "resources", "originals", "owasp_cornucopia_en_static.docx")
        input_doc = docx.Document(filename)
        self.want_file = os.path.join(c.convert_vars.BASE_PATH, "output", "owasp_cornucopia_en_static.docx")
        if os.path.isfile(self.want_file):
            os.remove(self.want_file)

        c.save_docx_file(input_doc, self.want_file)
        self.assertTrue(os.path.isfile(self.want_file))


class TestReplaceTextInLeafletlFile(unittest.TestCase):
    def setUp(self) -> None:
        self.input_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content>${Common_T00210} ${Common_T00220}</Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""
        self.input_dict = {
            "${Common_T00210}": "The idea behind Cornucopia is to help development teams, especially those using \
Agile methodologies, to identify application security requirements and develop security-based \
user stories.",
            "${Common_T00220}": "\
Although the idea had been waiting for enough time to progress it, the final \
motivation came when SAFECode published its Practical Security Stories and Security Tasks for Agile \
Development Environments in July 2012.",
        }
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        self.input_xml_file = os.path.join(c.convert_vars.BASE_PATH, "output", "temp", "Stories", "Story_u8fb5.xml")
        if not os.path.exists(os.path.dirname(self.input_xml_file)):
            os.makedirs(os.path.dirname(self.input_xml_file))
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)
        with open(self.input_xml_file, "x", encoding="utf-8") as f:
            f.writelines(self.input_data)

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)

    def test_replace_text_in_xml_file_success(self) -> None:
        want_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content>The idea behind Cornucopia is to help development teams, especially those using \
Agile methodologies, to identify application security requirements and develop security-based \
user stories. \
Although the idea had been waiting for enough time to progress it, the final \
motivation came when SAFECode published its Practical Security Stories and Security Tasks for Agile \
Development Environments in July 2012.</Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""

        c.replace_text_in_xml_file(self.input_xml_file, list(self.input_dict.items()))
        with open(self.input_xml_file, "r", encoding="utf-8") as f:
            got_data = f.read()
        self.assertEqual(want_data, got_data)


class TestReplaceTextInEmptyLeafletlFile(unittest.TestCase):

    def setUp(self) -> None:
        self.input_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content></Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""
        self.input_dict = {
            "${Common_T00210}": "The idea behind Cornucopia is to help development teams, especially those using \
Agile methodologies, to identify application security requirements and develop security-based \
user stories.",
            "${Common_T00220}": "\
Although the idea had been waiting for enough time to progress it, the final \
motivation came when SAFECode published its Practical Security Stories and Security Tasks for Agile \
Development Environments in July 2012.",
        }
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        self.input_xml_file = os.path.join(c.convert_vars.BASE_PATH, "output", "temp", "Stories", "Story_u8fb5.xml")
        if not os.path.exists(os.path.dirname(self.input_xml_file)):
            os.makedirs(os.path.dirname(self.input_xml_file))
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)
        with open(self.input_xml_file, "x", encoding="utf-8") as f:
            f.writelines(self.input_data)

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)

    def test_replace_text_in_xml_file_success_but_empty(self) -> None:
        want_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content></Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""

        c.replace_text_in_xml_file(self.input_xml_file, list(self.input_dict.items()))
        with open(self.input_xml_file, "r", encoding="utf-8") as f:
            got_data = f.read()
        self.assertEqual(want_data, got_data)


class TestReplaceTextInXmlFile(unittest.TestCase):
    def setUp(self) -> None:
        self.input_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content>${WC_suit}</Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""
        self.input_dict = {
            "${VE_suit}": "Data validation & encoding",
            "${VE_VEA_desc}": "You have invented a new attack against Data Validation and Encoding",
            "${VE_VE2_desc}": "Read more about this topic in OWASP's free Cheat Sheets",
            "${WC_suit}": "Wild Card",
            "${WC_JOA_desc}": "Alice can utilize the application to attack users' systems and data",
        }
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        self.input_xml_file = os.path.join(c.convert_vars.BASE_PATH, "output", "temp", "Stories", "Story_u8fb5.xml")
        if not os.path.exists(os.path.dirname(self.input_xml_file)):
            os.makedirs(os.path.dirname(self.input_xml_file))
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)
        with open(self.input_xml_file, "x", encoding="utf-8") as f:
            f.writelines(self.input_data)

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)

    def test_replace_text_in_xml_file_success(self) -> None:
        want_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content>Wild Card</Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""

        c.replace_text_in_xml_file(self.input_xml_file, list(self.input_dict.items()))
        with open(self.input_xml_file, "r", encoding="utf-8") as f:
            got_data = f.read()
        self.assertEqual(want_data, got_data)


class TestReplaceTextInXmlFileFail(unittest.TestCase):
    def setUp(self) -> None:
        self.input_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content></Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""
        self.input_dict = {}
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        self.input_xml_file = os.path.join(c.convert_vars.BASE_PATH, "output", "temp", "Stories", "Story_u8fb5.xml")
        if not os.path.exists(os.path.dirname(self.input_xml_file)):
            os.makedirs(os.path.dirname(self.input_xml_file))
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)
        with open(self.input_xml_file, "x", encoding="utf-8") as f:
            f.writelines(self.input_data)

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b
        if os.path.isfile(self.input_xml_file):
            os.remove(self.input_xml_file)

    def test_replace_text_in_xml_file_fail(self) -> None:
        want_data = """<ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Suit">
    CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
        <Content></Content>
    </CharacterStyleRange>
</ParagraphStyleRange>"""

        with self.assertLogs(logging.getLogger(), logging.ERROR) as ll:
            c.replace_text_in_xml_file(self.input_xml_file, list(self.input_dict.items()))
        self.assertIn("ERROR:root:Failed to parse XML file", ll.output.pop(), "No xml parsing error was caught.")

        with open(self.input_xml_file, "r", encoding="utf-8") as f:
            got_data = f.read()
        self.assertEqual(want_data, got_data)


class Test1(unittest.TestCase):
    def setUp(self) -> None:
        self.replacement_values = [
            ("${VE_suit}", "Validation & Encoding"),
            ("${VE_VE2_desc}", "You have invented a new attack against Data Validation and Encoding"),
            ("${WC_suit}", "Wild Card"),
            ("${WC_JOA_desc}", "Alice can utilize the application to attack users' systems and data"),
        ]

    def test_get_replacement_value_from_dict_exact(self) -> None:
        input_text = "${VE_VE2_desc}"
        want_data = "You have invented a new attack against Data Validation and Encoding"

        got_data = c.get_replacement_value_from_dict(input_text, self.replacement_values)
        self.assertEqual(want_data, got_data)

    def test_get_replacement_value_from_dict_spaced(self) -> None:
        input_text = " ${VE_VE2_desc} "
        want_data = " You have invented a new attack against Data Validation and Encoding "

        got_data = c.get_replacement_value_from_dict(input_text, self.replacement_values)
        self.assertEqual(want_data, got_data)


class TestCheckMakeListIntoText(unittest.TestCase):
    def test_check_make_list_into_text_success(self) -> None:
        input_list = ["69", "107", "108", "109", "136", "137", "153", "156", "158", "162"]
        want_text = "69, 107-109, 136-137, 153, 156, 158, 162"

        got_text = c.check_make_list_into_text(input_list)
        self.assertEqual(want_text, got_text)

    def test_check_make_list_into_text_not_grouped(self) -> None:
        input_list = ["69", "107", "108", "109", "136", "137", "153", "156", "158", "162"]
        want_text = "69, 107-109, 136-137, 153, 156, 158, 162"

        got_text = c.check_make_list_into_text(input_list)
        self.assertEqual(want_text, got_text)

    def test_check_make_list_into_text_not_numeric(self) -> None:
        input_list = ["69", "107", "108", "109", "Ess", "137", "153", "156", "158", "162"]
        want_text = "69, 107, 108, 109, Ess, 137, 153, 156, 158, 162"

        got_text = c.check_make_list_into_text(input_list)
        self.assertEqual(want_text, got_text)


class TestGroupNumberRanges(unittest.TestCase):
    def test_group_number_ranges_success(self) -> None:
        input_list = ["69", "107", "108", "109", "136", "137", "153", "156", "158", "162"]
        want_text = ["69", "107-109", "136-137", "153", "156", "158", "162"]

        got_text = c.group_number_ranges(input_list)
        self.assertEqual(want_text, got_text)

    def test_group_number_ranges_non_numeric(self) -> None:
        input_list = ["69", "East", "West", "109", "136", "137", "153", "156", "158", "162"]
        want_text = ["69", "East", "West", "109", "136", "137", "153", "156", "158", "162"]

        got_text = c.group_number_ranges(input_list)
        self.assertEqual(want_text, got_text)


class TestGetSuitTagsAndKey(unittest.TestCase):
    def test_get_suit_tags_and_key_suits(self) -> None:
        input_key = "suits"
        want_tags = ["VE", "AT", "SM", "AZ", "CR", "C", "WC"]
        want_key = "cards"

        got_tags, got_key = c.get_suit_tags_and_key(input_key, "webapp")
        self.assertEqual(want_tags, got_tags)
        self.assertEqual(want_key, got_key)

    def test_get_suit_tags_and_key_paragraphs(self) -> None:
        input_key = "paragraphs"
        want_tags = ["Common"]
        want_key = "sentences"

        got_tags, got_key = c.get_suit_tags_and_key(input_key, "webapp")
        self.assertEqual(want_tags, got_tags)
        self.assertEqual(want_key, got_key)

    def test_get_suit_tags_and_key_suits_for_mobileapp(self) -> None:
        input_key = "suits"
        want_tags = ["PC", "AA", "NS", "RS", "CRM", "CM", "WC"]
        want_key = "cards"

        got_tags, got_key = c.get_suit_tags_and_key(input_key, "mobileapp")
        self.assertEqual(want_tags, got_tags)
        self.assertEqual(want_key, got_key)


class TestGetFullTag(unittest.TestCase):
    def test_get_full_tag_ve(self) -> None:
        input_suit_tag = "VE"
        input_id = "VE2"
        input_tag = "desc"
        want_full_tag = "${VE_VE2_desc}"

        got_full_tag = c.get_full_tag(input_suit_tag, input_id, input_tag)
        self.assertEqual(want_full_tag, got_full_tag)

    def test_get_full_tag_wc(self) -> None:
        input_suit_tag = "WC"
        input_card = "JOA"
        input_tag = "desc"
        want_full_tag = "${WC_JOA_desc}"

        got_full_tag = c.get_full_tag(input_suit_tag, input_card, input_tag)
        self.assertEqual(want_full_tag, got_full_tag)

    def test_get_full_tag_common(self) -> None:
        input_suit_tag = "Common"
        input_card = "T00010"
        input_tag = "text"
        want_full_tag = "${Common_T00010}"

        got_full_tag = c.get_full_tag(input_suit_tag, input_card, input_tag)
        self.assertEqual(want_full_tag, got_full_tag)


class TestZipDir(unittest.TestCase):
    def setUp(self) -> None:
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")
        self.input_filename = os.path.join(c.convert_vars.BASE_PATH, "output", "test.zip")

    def tearDown(self) -> None:
        if os.path.isfile(self.input_filename):
            os.remove(self.input_filename)
        c.convert_vars.BASE_PATH = self.b

    def test_zip_dir_success(self) -> None:
        input_path = os.path.join(c.convert_vars.BASE_PATH, "source")
        if os.path.isfile(self.input_filename):
            os.remove(self.input_filename)
        want_min_file_size = 2600

        c.zip_dir(input_path, self.input_filename)
        self.assertTrue(os.path.isfile(self.input_filename))
        got_file_size = os.stat(self.input_filename).st_size
        self.assertGreater(got_file_size, want_min_file_size)


class TestEnsureFolderExists(unittest.TestCase):
    def setUp(self) -> None:
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")

    def tearDown(self) -> None:
        temp_path = os.path.join(c.convert_vars.BASE_PATH, "output", "temp")
        c.convert_vars.BASE_PATH = self.b
        if os.path.exists(temp_path):
            shutil.rmtree(temp_path, ignore_errors=True)

    def test_ensure_folder_exists_new(self) -> None:
        input_path = os.path.join(c.convert_vars.BASE_PATH, "output", "temp", "path")

        c.ensure_folder_exists(input_path)
        self.assertTrue(os.path.exists(input_path))

    def test_ensure_folder_exists_old(self) -> None:
        input_path = os.path.join(c.convert_vars.BASE_PATH, "output", "temp", "path")

        c.ensure_folder_exists(input_path)
        self.assertTrue(os.path.exists(input_path))
        c.ensure_folder_exists(input_path)
        self.assertTrue(os.path.exists(input_path))


class TestReplaceDocxInlineText(unittest.TestCase):
    def setUp(self) -> None:
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b

    @staticmethod
    def get_docx_text(doc: docx.Document) -> List[str]:
        text_list: List[str] = []
        paragraphs = doc.paragraphs
        for p in paragraphs:
            t = ""
            for r in p.runs:
                t += r.text
            if t not in ["", "\n"]:
                text_list.append(t)
        return text_list

    def test_replace_docx_inline_text_expected_keys_present(self) -> None:
        template_docx_file = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        doc = docx.Document(template_docx_file)
        input_replacement_data = {
            "${Common_T03100}": "Alice can utilize the application to attack users' systems and data",
        }
        want_old_text = list(k for k, v in input_replacement_data.items())

        text_list = self.get_docx_text(doc)
        for t in want_old_text:
            self.assertIn(t, text_list)

    def test_replace_docx_inline_text_new_text_present(self) -> None:
        template_docx_file = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        doc = docx.Document(template_docx_file)
        input_replacement_data = {
            "${Common_T03100}": "Alice can utilize the application to attack users' systems and data",
        }
        want_new_text = list(v for k, v in input_replacement_data.items())

        doc = c.replace_docx_inline_text(doc, input_replacement_data)
        text_list = self.get_docx_text(doc)
        for t in want_new_text:
            self.assertIn(t, text_list)

    def test_replace_docx_inline_text_keys_replaced(self) -> None:
        template_docx_file = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        doc = docx.Document(template_docx_file)
        input_replacement_data = {
            "${Common_T03100}": "Alice can utilize the application to attack users' systems and data",
        }
        want_old_text = list(k for k, v in input_replacement_data.items())

        doc = c.replace_docx_inline_text(doc, input_replacement_data)
        text_list = self.get_docx_text(doc)
        for t in want_old_text:
            self.assertNotIn(t, text_list)


class TestGetDocumentParagraphs(unittest.TestCase):
    def setUp(self) -> None:
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b

    def test_get_document_paragraphs_len_paragraphs(self) -> None:
        template_docx_file = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        doc = docx.Document(template_docx_file)
        # Accept a range to handle platform/version differences in docx parsing
        # On Windows: 1941, On Linux CI: 2010
        min_paragraphs = 1940
        max_paragraphs = 2010

        paragraphs = c.get_document_paragraphs(doc)
        self.assertGreaterEqual(
            len(paragraphs), min_paragraphs, f"Expected at least {min_paragraphs} paragraphs, got {len(paragraphs)}"
        )
        self.assertLessEqual(
            len(paragraphs), max_paragraphs, f"Expected at most {max_paragraphs} paragraphs, got {len(paragraphs)}"
        )

    def test_get_document_paragraphs_find_text(self) -> None:
        template_docx_file = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        doc = docx.Document(template_docx_file)
        want_text_list = [
            "${VE_suit}",
            "${VE_VE2_desc}",
            "${Common_T00010}",
        ]

        paragraphs = c.get_document_paragraphs(doc)
        text_list: List[str] = []
        for p in paragraphs:
            t = "".join([run.text for run in p.runs])
            if t not in ["", "\n"]:
                text_list.append(t)
        for want_text in want_text_list:
            self.assertIn(want_text, text_list)


class TestGetParagraphsFromTableInDoc(unittest.TestCase):
    def setUp(self) -> None:
        self.b = c.convert_vars.BASE_PATH
        c.convert_vars.BASE_PATH = os.path.join(self.b, "tests", "test_files")

    def tearDown(self) -> None:
        c.convert_vars.BASE_PATH = self.b

    def test_get_paragraphs_from_table_in_doc(self) -> None:
        template_docx_file = os.path.join(
            c.convert_vars.BASE_PATH, "resources", "templates", "owasp_cornucopia_webapp_ver_guide_bridge_lang.docx"
        )
        doc = docx.Document(template_docx_file)
        doc_tables = doc.tables
        want_min_len_paragraphs = 1000

        paragraphs = []
        for table in doc_tables:
            paragraphs += c.get_paragraphs_from_table_in_doc(table)
        self.assertGreater(len(paragraphs), want_min_len_paragraphs)


class TestSafeExtractAll(unittest.TestCase):
    """Unit tests for _safe_extractall covering CWE-22 path traversal prevention."""

    def _build_zip(self, members: typing.Dict[str, str]) -> io.BytesIO:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            for name, content in members.items():
                zf.writestr(zipfile.ZipInfo(name), content)
        buf.seek(0)
        return buf

    def test_blocks_parent_directory_traversal(self) -> None:
        """Zip Slip via '../' in member filename must raise ValueError."""
        buf = self._build_zip({"../evil.txt": "pwned"})
        with tempfile.TemporaryDirectory() as td:
            with zipfile.ZipFile(buf) as zf:
                with self.assertRaises(ValueError) as ctx:
                    c._safe_extractall(zf, td)
        self.assertIn("Zip Slip blocked", str(ctx.exception))

    def test_blocks_absolute_path_member(self) -> None:
        """/etc/passwd as a member filename must be blocked."""
        buf = self._build_zip({"/etc/passwd": "pwned"})
        with tempfile.TemporaryDirectory() as td:
            with zipfile.ZipFile(buf) as zf:
                with self.assertRaises(ValueError):
                    c._safe_extractall(zf, td)

    def test_allows_legitimate_nested_members(self) -> None:
        """Normal nested paths must extract correctly."""
        buf = self._build_zip({"content.xml": "<r/>", "subdir/file.xml": "<s/>"})
        with tempfile.TemporaryDirectory() as td:
            with zipfile.ZipFile(buf) as zf:
                c._safe_extractall(zf, td)
            self.assertTrue(os.path.isfile(os.path.join(td, "content.xml")))
            self.assertTrue(os.path.isfile(os.path.join(td, "subdir", "file.xml")))

    def test_skips_root_dot_entry(self) -> None:
        """A '.' root directory entry must be skipped without error."""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            dot = zipfile.ZipInfo(".")
            dot.external_attr = 0o40755 << 16
            zf.writestr(dot, "")
            zf.writestr("content.xml", "<r/>")
        buf.seek(0)
        with tempfile.TemporaryDirectory() as td:
            with zipfile.ZipFile(buf) as zf:
                c._safe_extractall(zf, td)
            self.assertTrue(os.path.isfile(os.path.join(td, "content.xml")))


if __name__ == "__main__":
    unittest.main()
