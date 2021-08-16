import unittest
import argparse
import os
import docx  # type: ignore
import logging
import typing
from typing import List, Dict, Any, Tuple

import scripts.convert as c

c.convert_vars = c.ConvertVars()


class TestGetValidFileTypes(unittest.TestCase):
    def test_get_valid_file_types_idml(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfiletype="idml")
        want_list = ["idml"]

        got_list = c.get_valid_file_types()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_blank_filetype(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfiletype="", outputfile="cornucopia_en_template.pdf")
        want_list = ["pdf"]

        got_list = c.get_valid_file_types()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_revert_default(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfiletype="", outputfile="cornucopia_en_template")
        want_list = ["docx"]

        got_list = c.get_valid_file_types()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_all_with_pdf(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfiletype="all")
        c.convert_vars.can_convert_to_pdf = True
        want_list = ["docx", "idml", "pdf"]

        got_list = c.get_valid_file_types()
        got_list.sort()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_all_without_pdf(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfiletype="all")
        c.convert_vars.can_convert_to_pdf = False
        want_list = ["docx", "idml"]

        got_list = c.get_valid_file_types()
        got_list.sort()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_pdf_without_pdf_ability(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfiletype="pdf")
        c.convert_vars.can_convert_to_pdf = False
        want_list: List[str] = []

        logging.getLogger().setLevel(logging.CRITICAL)
        got_list = c.get_valid_file_types()
        logging.getLogger().setLevel(logging.ERROR)
        self.assertListEqual(want_list, got_list)


class TestGetValidLanguagesChoices(unittest.TestCase):
    def test_get_valid_language_choices_fr(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="fr")
        want_language = ["fr"]

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_blank(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="")
        want_language = ["en"]

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_all(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="all")
        want_language = c.convert_vars.LANGUAGE_CHOICES
        want_language.remove("template")
        want_language.remove("all")

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_template(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="template")
        want_language = ["template"]

        got_language = c.get_valid_language_choices()
        self.assertListEqual(want_language, got_language)


class TestSetCanConvertToPdf(unittest.TestCase):
    def test_set_can_convert_to_pdf(self) -> None:
        # c.convert_vars.can_convert_to_pdf = None
        want_can_convert_in: List[bool] = [True, False]

        c.set_can_convert_to_pdf()
        got_can_convert = c.convert_vars.can_convert_to_pdf
        self.assertIn(got_can_convert, want_can_convert_in)


class TestSortKeysLongestToShortest(unittest.TestCase):
    def test_sort_keys_longest_to_shortest_success(self) -> None:
        source_data = {
            "key shortest": "value1",
            "key ------------- longest": "value2",
            "key ------ middle": "value3",
        }
        want_data = [
            ("key ------------- longest", "value2"),
            ("key ------ middle", "value3"),
            ("key shortest", "value1"),
        ]

        got_data = c.sort_keys_longest_to_shortest(source_data)
        self.assertEqual(want_data, got_data)


class TestSetLogging(unittest.TestCase):
    def test_set_logging_level_default(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        want_logging_level = logging.INFO

        c.set_logging()
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)

    def test_set_logging_level_debug(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=True)
        want_logging_level = logging.DEBUG

        c.set_logging()
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)

    def test_set_logging_level_info(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=None)
        want_logging_level = logging.INFO

        c.set_logging()
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)


class TestRemoveShortKeys(unittest.TestCase):
    def test_remove_short_keys_10_chars(self) -> None:
        input_dict = {
            "1": "txt2",
            "123": "txt123",
            "0123456789": "txt0123456789",
            "0123456789012345678901234567890123456789": "txt0123456789",
        }
        min_length = 10
        want_dict = {
            "0123456789": "txt0123456789",
            "0123456789012345678901234567890123456789": "txt0123456789",
        }

        got_dict = c.remove_short_keys(input_dict, min_length)
        self.assertDictEqual(want_dict, got_dict)

    def test_remove_short_keys_default_40_chars(self) -> None:
        input_dict = {
            "1": "txt2",
            "123": "txt123",
            "0123456789": "txt0123456789",
            "01234567890123456789012345678901234567890123456789012345678901234567890123456789": "txt long",
        }
        want_dict = {
            "01234567890123456789012345678901234567890123456789012345678901234567890123456789": "txt long",
        }

        got_dict = c.remove_short_keys(input_dict)
        self.assertDictEqual(want_dict, got_dict)


class TestGetTemplateDoc(unittest.TestCase):
    def test_get_template_doc_default_docx(self) -> None:
        c.convert_vars.args = argparse.Namespace(inputfile="")
        c.convert_vars.making_template = False
        input_filetype = "docx"
        want_template_doc = (
            c.convert_vars.BASE_PATH + "/resources/templates/owasp_cornucopia_edition_lang_ver_template.docx"
        )

        got_template_doc = c.get_template_doc(input_filetype)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_doc_default_idml(self) -> None:
        c.convert_vars.args = argparse.Namespace(inputfile="")
        c.convert_vars.making_template = False
        input_filetype = "idml"
        want_template_doc = (
            c.convert_vars.BASE_PATH + "/resources/templates/owasp_cornucopia_edition_lang_ver_template.idml"
        )

        got_template_doc = c.get_template_doc(input_filetype)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_doc_make_template_docx(self) -> None:
        c.convert_vars.args = argparse.Namespace(inputfile="")
        c.convert_vars.making_template = True
        input_filetype = "docx"
        want_template_doc = c.convert_vars.BASE_PATH + "/resources/originals/owasp_cornucopia_en.docx"

        got_template_doc = c.get_template_doc(input_filetype)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_doc_make_template_idml(self) -> None:
        c.convert_vars.args = argparse.Namespace(inputfile="")
        c.convert_vars.making_template = True
        input_filetype = "idml"
        want_template_doc = c.convert_vars.BASE_PATH + "/resources/originals/owasp_cornucopia_en.idml"

        got_template_doc = c.get_template_doc(input_filetype)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_doc_relative_path(self) -> None:
        c.convert_vars.args = argparse.Namespace(
            inputfile="../test/test_files/owasp_cornucopia_edition_lang_ver_template.docx"
        )
        c.convert_vars.making_template = False
        input_filetype = "docx"
        want_template_doc = (
            c.convert_vars.BASE_PATH + "/test/test_files/owasp_cornucopia_edition_lang_ver_template.docx"
        )

        got_template_doc = c.get_template_doc(input_filetype)
        self.assertEqual(want_template_doc, got_template_doc)

    def test_get_template_doc_file_not_exist(self) -> None:
        c.convert_vars.args = argparse.Namespace(inputfile="../resources/templates/owasp_cornucopia_template.docx")
        c.convert_vars.making_template = False
        input_filetype = "docx"
        want_template_doc = ""

        logging.getLogger().setLevel(logging.CRITICAL)
        got_template_doc = c.get_template_doc(input_filetype)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertEqual(want_template_doc, got_template_doc)


class TestRenameOutputFile(unittest.TestCase):
    def test_rename_output_file_short(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfile="/output/cornucopia_edition_component_lang_ver.docx")
        input_file_type = "docx"
        input_meta_data = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "121"}
        want_filename = "/output/cornucopia_ecommerce_cards_en_121.docx"

        got_filename = c.rename_output_file(input_file_type, input_meta_data)
        self.assertEqual(want_filename, got_filename)

    def test_rename_output_file_no_extension(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfile="/output/cornucopia_edition_component_lang_ver")
        input_file_type = "idml"
        input_meta_data = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "121"}
        want_filename = "/output/cornucopia_ecommerce_cards_en_121.idml"

        got_filename = c.rename_output_file(input_file_type, input_meta_data)
        self.assertEqual(want_filename, got_filename)

    def test_rename_output_file_using_defaults(self) -> None:
        c.convert_vars.args = argparse.Namespace(outputfile=c.convert_vars.DEFAULT_OUTPUT_FILENAME)
        input_file_type = "docx"
        input_meta_data = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"}
        want_filename = c.convert_vars.BASE_PATH + "/output/owasp_cornucopia_ecommerce_cards_en_1.21.docx"

        # logging.getLogger().setLevel(logging.DEBUG)
        got_filename = c.rename_output_file(input_file_type, input_meta_data)
        self.assertEqual(want_filename, got_filename)


class TestGetFindReplaceList(unittest.TestCase):
    want_list_default: List[Tuple[str, str]] = [
        ("_type", "_ecommerce"),
        ("_edition", "_ecommerce"),
        ("_component", "_cards"),
        ("_language", "_en"),
        ("_lang", "_en"),
        ("_version", "_1.21"),
        ("_ver", "_1.21"),
    ]

    def test_get_find_replace_list_default(self) -> None:
        c.convert_vars.making_template = False
        input_meta_data = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"}
        want_list = self.want_list_default.copy()
        want_list.append(("_template", ""))

        got_list = c.get_find_replace_list(input_meta_data)
        self.assertListEqual(want_list, got_list)

    def test_get_find_replace_list_making_template(self) -> None:
        c.convert_vars.making_template = True
        input_meta_data = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"}
        want_list = self.want_list_default.copy()

        got_list = c.get_find_replace_list(input_meta_data)
        self.assertListEqual(want_list, got_list)


class TestGetMetaData(unittest.TestCase):
    test_data: Dict[str, Any] = {
        "meta": {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"},
        "suits": [
            {
                "name": "Data validation & encoding",
                "cards": [
                    {
                        "value": "A",
                        "desc": "You have invented a new attack against Data Validation",
                        "misc": "Read more about this topic in OWASP's free Cheat Sheets",
                    },
                    {"value": "2", "desc": "Brian can gather information about the underlying configurations"},
                ],
            },
        ],
    }

    def test_get_meta_data_defaults(self) -> None:
        input_data = self.test_data.copy()
        want_data = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"}

        got_data = c.get_meta_data(input_data)
        self.assertDictEqual(want_data, got_data)

    def test_get_meta_data_failure(self) -> None:
        input_data = self.test_data.copy()
        del input_data["meta"]
        want_data: Dict[str, str] = {}

        logging.getLogger().setLevel(logging.CRITICAL)
        got_data = c.get_meta_data(input_data)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertDictEqual(want_data, got_data)




class TestGetReplacementData(unittest.TestCase):
    test_data: Dict[str, Any] = {
        "meta": {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"},
        "suits": [
            {
                "name": "Data validation & encoding",
                "cards": [
                    {
                        "value": "A",
                        "desc": "You have invented a new attack against Data Validation and Encoding",
                        "misc": "Read more about this topic in OWASP's free Cheat Sheets on Input Validation, "
                        "XSS Prevention, DOM-based XSS Prevention, SQL Injection Prevention, "
                        "and Query Parameterization",
                    },
                    {"value": "2", "desc": "Brian can gather information about the underlying configurations"},
                ],
            },
        ],
        "paragraphs": [
            {
                "name": "Common",
                "cards": [
                    {
                        "value": "NoCard",
                        "text": "No Card",
                    },
                    {
                        "value": "Title",
                        "text": "OWASP Cornucopia Ecommerce Edition v1.21-EN",
                    },
                ],
            },
        ],
    }

    def test_get_replacement_data_translation_meta(self) -> None:
        input_yaml_files = [
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-cards-1.21-en.yaml",
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-mappings-1.2.yaml",
        ]
        input_data_type = "translation"
        input_language = "en"
        want_meta = {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"}

        got_data = c.get_replacement_data(input_yaml_files, input_data_type, input_language)
        self.assertEqual(want_meta, got_data["meta"])

    def test_get_replacement_data_translation_first_suit_first_card(self) -> None:
        input_yaml_files = [
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-cards-1.21-en.yaml",
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-mappings-1.2.yaml",
        ]
        input_data_type = "translation"
        input_language = "en"
        want_first_suit_keys = self.test_data["suits"][0].keys()
        want_first_suit_first_card_keys = self.test_data["suits"][0]["cards"][0].keys()
        want_first_suit_first_card_value = self.test_data["suits"][0]["cards"][0]["value"]

        got_suits = c.get_replacement_data(input_yaml_files, input_data_type, input_language)["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card_keys = got_suits[0]["cards"][0].keys()
        self.assertEqual(want_first_suit_first_card_keys, got_first_suit_first_card_keys)
        got_first_suit_first_card_value = got_suits[0]["cards"][0]["value"]
        self.assertEqual(want_first_suit_first_card_value, got_first_suit_first_card_value)

    def test_get_replacement_data_mappings_meta(self) -> None:
        input_yaml_files = [
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-cards-1.21-en.yaml",
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-mappings-1.2.yaml",
        ]
        input_data_type = "mappings"
        input_language = "en"
        want_meta = {"edition": "ecommerce", "component": "mappings", "language": "ALL", "version": "1.2"}

        got_data = c.get_replacement_data(input_yaml_files, input_data_type, input_language)
        self.assertEqual(want_meta, got_data["meta"])

    def test_get_replacement_data_mappings_first_suit_first_card(self) -> None:
        input_yaml_files = [
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-cards-1.21-en.yaml",
            c.convert_vars.BASE_PATH + "/test/test_files/ecommerce-mappings-1.2.yaml",
        ]
        input_data_type = "mappings"
        input_language = "en"
        want_first_suit_keys = {"name": "", "cards": ""}.keys()
        want_first_suit_first_card = {
            "value": "2",
            "owasp_scp": [69, 107, 108, 109, 136, 137, 153, 156, 158, 162],
            "owasp_asvs": [1.10, 4.5, 8.1, 11.5, 19.1, 19.5],
            "owasp_appsensor": ["HT1", "HT2", "HT3"],
            "capec": [54, 541],
            "safecode": [4, 23],
        }

        got_suits = c.get_replacement_data(input_yaml_files, input_data_type, input_language)["suits"]
        got_first_suit_keys = got_suits[0].keys()
        self.assertEqual(want_first_suit_keys, got_first_suit_keys)
        got_first_suit_first_card = got_suits[0]["cards"][0]
        self.assertEqual(want_first_suit_first_card, got_first_suit_first_card)


class TestParseArguments(unittest.TestCase):
    def test_successfully_parsing_arguments_short_form_basic(self) -> None:
        input_args = ["-t", "idml"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="en",
            debug=False,
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)

    def test_successfully_parsing_arguments_short_form_language(self) -> None:
        input_args = ["-t", "idml", "-l", "fr"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="fr",
            debug=False,
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)

    def test_successfully_parsing_arguments_long_form_basic(self) -> None:
        input_args = ["--outputfiletype", "idml"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="en",
            debug=False,
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)

    def test_successfully_parsing_arguments_long_form_language(self) -> None:
        input_args = ["--outputfiletype", "idml", "--language", "fr"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="fr",
            debug=False,
        )

        got_args = c.parse_arguments(input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)


class TestMakeTemplate(unittest.TestCase):
    def test_set_making_template_true(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="template")

        c.set_making_template()
        result = c.convert_vars.making_template
        self.assertTrue(result)

    def test_set_making_template_false(self) -> None:
        c.convert_vars.args = argparse.Namespace(language="en")

        c.set_making_template()
        result = c.convert_vars.making_template
        self.assertFalse(result)

    def test_set_making_template_empty(self) -> None:
        c.convert_vars.args = argparse.Namespace()

        c.set_making_template()
        result = c.convert_vars.making_template
        self.assertFalse(result)


class TestGetFilesFromOfType(unittest.TestCase):
    def test_get_files_from_of_type_source_yaml_files(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        path = c.convert_vars.BASE_PATH + "/test/test_files/"
        ext = "yaml"
        want_files = ["ecommerce-cards-1.21-en.yaml", "ecommerce-mappings-1.2.yaml"]
        want_count = len(want_files)

        got_files = c.get_files_from_of_type(path, ext)
        self.assertEqual(len(got_files), want_count)
        got_files = list(os.path.basename(f) for f in got_files)
        got_files.sort()
        self.assertListEqual(got_files, want_files)

    def test_get_files_from_of_type_source_docx_files(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        path = c.convert_vars.BASE_PATH + "/test/test_files/"
        ext = "docx"
        want_files = ["owasp_cornucopia_edition_lang_ver_template.docx"]

        got_files = c.get_files_from_of_type(path, ext)
        got_files = list(os.path.basename(f) for f in got_files)
        self.assertListEqual(got_files, want_files)
        for f in want_files:
            self.assertIn(f, got_files)

    def test_get_files_from_of_type_source_empty_list(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        path = c.convert_vars.BASE_PATH + "/test/test_files/"
        ext = "ext"
        want_files: typing.List[str] = []

        logging.getLogger().setLevel(logging.CRITICAL)
        got_files = c.get_files_from_of_type(path, ext)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertListEqual(got_files, want_files)


class TestGetDocxDocument(unittest.TestCase):
    def test_get_docx_document_success(self) -> None:
        file = c.convert_vars.BASE_PATH + "/test/test_files/owasp_cornucopia_edition_lang_ver_template.docx"
        want_type = docx.document.Document
        want_len_paragraphs = 36

        got_file = c.get_docx_document(file)
        got_type = type(got_file)
        self.assertEqual(want_type, got_type)
        got_len_paragraphs = len(got_file.paragraphs)
        self.assertEqual(want_len_paragraphs, got_len_paragraphs)

    def test_get_docx_document_failure(self) -> None:
        file = c.convert_vars.BASE_PATH + "/test/test_files/owasp_cornucopia_edition_lang_ver_template.d"
        want_type = docx.document.Document
        want_len_paragraphs = 0

        logging.getLogger().setLevel(logging.CRITICAL)
        got_file = c.get_docx_document(file)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertIsInstance(got_file, want_type)
        got_len_paragraphs = len(got_file.paragraphs)
        self.assertEqual(want_len_paragraphs, got_len_paragraphs)


class TestGetTagForSuitName(unittest.TestCase):
    def test_get_tag_for_suit_name_ve(self) -> None:
        c.convert_vars.making_template = False
        suit = {"name": "Data validation & encoding", "cards": []}
        suit_tag = "VE"
        want_tag_data = {"${VE_suit}": suit["name"]}

        got_tag_data = c.get_tag_for_suit_name(suit, suit_tag)
        self.assertDictEqual(want_tag_data, got_tag_data)

    def test_get_tag_for_suit_name_ve_template(self) -> None:
        c.convert_vars.making_template = True
        suit = {"name": "Data validation & encoding", "cards": []}
        suit_tag = "VE"
        want_tag_data = {suit["name"]: "${VE_suit}"}

        # logging.basicConfig(level=logging.DEBUG)
        got_tag_data = c.get_tag_for_suit_name(suit, suit_tag)
        logging.basicConfig(level=logging.ERROR)
        self.assertDictEqual(want_tag_data, got_tag_data)


class TestGetReplacementDict(unittest.TestCase):
    test_data = {
        "meta": {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"},
        "suits": [
            {
                "name": "Data validation & encoding",
                "cards": [
                    {
                        "value": "A",
                        "desc": "You have invented a new attack against Data Validation",
                        "misc": "Read more about this topic in OWASP's free Cheat Sheets",
                    },
                    {"value": "2", "desc": "Brian can gather information about the underlying configurations"},
                ],
            },
            {
                "name": "Authentication",
                "cards": [
                    {
                        "value": "A",
                        "desc": "You have invented a new attack against Authentication",
                        "misc": "Read more about this topic in OWASP's free Cheat Sheet",
                    },
                    {"value": "2", "desc": "James can undertake authentication functions without"},
                ],
            },
        ],
    }

    def test_get_replacement_dict_success(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        c.convert_vars.making_template = False
        logging.basicConfig(level=logging.ERROR)
        input_data = self.test_data.copy()
        want_type = dict
        want_length = 8
        want_data = {
            "${VE_suit}": "Data validation & encoding",
            "${VE_VEA_desc}": "You have invented a new attack against Data Validation",
            "${VE_VEA_misc}": "Read more about this topic in OWASP's free Cheat Sheets",
            "${VE_VE2_desc}": "Brian can gather information about the underlying configurations",
            "${AT_suit}": "Authentication",
            "${AT_ATA_desc}": "You have invented a new attack against Authentication",
            "${AT_ATA_misc}": "Read more about this topic in OWASP's free Cheat Sheet",
            "${AT_AT2_desc}": "James can undertake authentication functions without",
        }

        got_data = c.get_replacement_dict(input_data)
        self.assertIsInstance(got_data, want_type)
        self.assertEqual(len(got_data), want_length)
        self.assertDictEqual(got_data, want_data)

    def test_get_replacement_dict_template(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        c.convert_vars.making_template = True
        logging.basicConfig(level=logging.ERROR)
        input_data = self.test_data.copy()
        want_type = dict
        want_length = 8
        want_data = {
            "Data validation & encoding": "${VE_suit}",
            "You have invented a new attack against Data Validation": "${VE_VEA_desc}",
            "Read more about this topic in OWASP's free Cheat Sheets": "${VE_VEA_misc}",
            "Brian can gather information about the underlying configurations": "${VE_VE2_desc}",
            "Authentication": "${AT_suit}",
            "You have invented a new attack against Authentication": "${AT_ATA_desc}",
            "Read more about this topic in OWASP's free Cheat Sheet": "${AT_ATA_misc}",
            "James can undertake authentication functions without": "${AT_AT2_desc}",
        }

        got_data = c.get_replacement_dict(input_data=input_data)
        self.assertIsInstance(got_data, want_type)
        self.assertEqual(len(got_data), want_length)
        self.assertDictEqual(got_data, want_data)


class TestGetCheckFixFileExtension(unittest.TestCase):
    def test_get_check_fix_file_extension_no_extension(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        input_filename = "hello"
        input_extension = "docx"
        want_filename = "hello.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_no_extension_with_version(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        input_filename = "hello_v1.21"
        input_extension = "docx"
        want_filename = "hello_v1.21.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_with_leading_dot(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        input_filename = "hello"
        input_extension = ".docx"
        want_filename = "hello.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_wrong_extension(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        input_filename = "hello.docx"
        input_extension = "pdf"
        want_filename = "hello.pdf"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_correct_extension(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        input_filename = "hello.docx"
        input_extension = ".docx"
        want_filename = "hello.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_with_folders(self) -> None:
        c.convert_vars.args = argparse.Namespace(debug=False)
        input_filename = "output/folder/hello_v1.21"
        input_extension = "docx"
        want_filename = "output/folder/hello_v1.21.docx"

        got_filename = c.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

# class TestSaveIdmlFile
class TestWritePdfFile:
    def test_write_pdf_file_true(self) -> None:



    def test_write_pdf_file_false(self) -> None:
        want_data = ""





class TestGetMappingDict(unittest.TestCase):
    def test_get_mapping_dict_true(self, yaml_files=None) -> None:
        input_filename = yaml_files
        want_mapping_data = c.get_mapping_dict(yaml_files)
        got_filename = c.get_mapping_dict(input_filename, want_mapping_data)
        self.assertEqual(got_filename)

    def test_get_mapping_dict_false(self, yaml_files=None) -> None:
        input_filename = yaml_files
        want_mapping_data = ""

        got_filename = c.get_mapping_dict(input_filename, want_mapping_data)
        self.assertEqual(got_filename)
